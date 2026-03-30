"""
Generate a comprehensive Word document with:
1. Original comparison (Ollama self-judged vs Groq self-judged)
2. Fair comparison (both judged by Groq 70B)
3. All charts embedded
4. Detailed analysis of why results differ
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

C_OLLAMA = "#4A90D9"
C_GROQ = "#E8833A"
C_OLLAMA_FAIR = "#7BB3E0"
C_GROQ_FAIR = "#F0A86A"

QUESTIONS_SHORT = ["Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention", "Type 1\nDiabetes", "MRSA"]
QUESTIONS_FULL = [
    "Symptoms of bacterial vaginosis",
    "Causes of varicose veins",
    "Preventing melanoma",
    "Living with type 1 diabetes",
    "MRSA",
]

# ---- Original scores (self-judged) ----
orig_ollama = {
    "faith": [0.80, 0.80, 0.33, 0.50, 0.40],
    "recall": [0.40, 0.40, 0.40, 0.33, 0.40],
    "bleu": [0.911, 0.459, 0.785, 0.096, 0.342],
}
orig_groq = {
    "faith": [1.0, 1.0, 1.0, 1.0, 1.0],
    "recall": [1.0, 1.0, 1.0, 1.0, 1.0],
    "bleu": [0.868, 0.361, 0.969, 0.200, 0.786],
}

# ---- Fair scores (both judged by Groq 70B) ----
fair_ollama = {
    "faith": [1.0, 1.0, 1.0, 0.8, 0.6],
    "recall": [1.0, 1.0, 1.0, 1.0, 1.0],
    "bleu": [0.911, 0.459, 0.785, 0.096, 0.342],
}
fair_groq = {
    "faith": [1.0, 1.0, 1.0, 1.0, 1.0],
    "recall": [1.0, 1.0, 1.0, 1.0, 1.0],
    "bleu": [0.868, 0.361, 0.969, 0.200, 0.786],
}

avg = lambda lst: sum(lst) / len(lst)

# =========================================================
# GENERATE ALL CHARTS
# =========================================================
def make_bar_chart(filename, title, labels, ollama_vals, groq_vals, ol_label="Ollama (llama3.2)", gq_label="Groq (llama-3.3-70b)", wide=False):
    fig, ax = plt.subplots(figsize=(10, 5) if wide else (8, 5))
    x = np.arange(len(labels))
    w = 0.32
    b1 = ax.bar(x - w/2, [v*100 for v in ollama_vals], w, label=ol_label, color=C_OLLAMA)
    b2 = ax.bar(x + w/2, [v*100 for v in groq_vals], w, label=gq_label, color=C_GROQ)
    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title(title, fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9 if wide else 11)
    ax.set_ylim(0, 118)
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3)
    for b in b1:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=8)
    for b in b2:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=8)
    plt.tight_layout()
    plt.savefig(filename, dpi=150)
    plt.close()

def make_grouped_4bar(filename, title, labels, o_orig, g_orig, o_fair, g_fair):
    """4 bars per question: original ollama, original groq, fair ollama, fair groq"""
    fig, ax = plt.subplots(figsize=(12, 5.5))
    x = np.arange(len(labels))
    w = 0.18
    ax.bar(x - 1.5*w, [v*100 for v in o_orig], w, label="Ollama (self-judged)", color=C_OLLAMA)
    ax.bar(x - 0.5*w, [v*100 for v in g_orig], w, label="Groq (self-judged)", color=C_GROQ)
    ax.bar(x + 0.5*w, [v*100 for v in o_fair], w, label="Ollama (Groq-judged)", color=C_OLLAMA_FAIR, edgecolor=C_OLLAMA, linewidth=1.2)
    ax.bar(x + 1.5*w, [v*100 for v in g_fair], w, label="Groq (Groq-judged)", color=C_GROQ_FAIR, edgecolor=C_GROQ, linewidth=1.2)
    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title(title, fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylim(0, 118)
    ax.legend(fontsize=9, loc="upper right")
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(filename, dpi=150)
    plt.close()

# Original comparison charts
make_bar_chart("rpt_orig_overall.png", "Original Evaluation: Overall Averages",
    ["Faithfulness", "Context Recall", "BLEU Score"],
    [avg(orig_ollama["faith"]), avg(orig_ollama["recall"]), avg(orig_ollama["bleu"])],
    [avg(orig_groq["faith"]), avg(orig_groq["recall"]), avg(orig_groq["bleu"])])

for key, label in [("faith", "Faithfulness"), ("recall", "Context Recall"), ("bleu", "BLEU Score")]:
    make_bar_chart(f"rpt_orig_{key}.png", f"Original: {label} per Question",
        QUESTIONS_SHORT, orig_ollama[key], orig_groq[key], wide=True)

# Fair comparison charts
make_bar_chart("rpt_fair_overall.png", "Fair Comparison: Overall Averages (Same Groq 70B Judge)",
    ["Faithfulness", "Context Recall", "BLEU Score"],
    [avg(fair_ollama["faith"]), avg(fair_ollama["recall"]), avg(fair_ollama["bleu"])],
    [avg(fair_groq["faith"]), avg(fair_groq["recall"]), avg(fair_groq["bleu"])],
    ol_label="Ollama answers (Groq-judged)", gq_label="Groq answers (Groq-judged)")

for key, label in [("faith", "Faithfulness"), ("recall", "Context Recall"), ("bleu", "BLEU Score")]:
    make_bar_chart(f"rpt_fair_{key}.png", f"Fair: {label} per Question (Same Judge)",
        QUESTIONS_SHORT, fair_ollama[key], fair_groq[key], wide=True,
        ol_label="Ollama answers", gq_label="Groq answers")

# Side-by-side original vs fair
make_grouped_4bar("rpt_sidebyside_faith.png", "Faithfulness: Original vs Fair Comparison",
    QUESTIONS_SHORT, orig_ollama["faith"], orig_groq["faith"], fair_ollama["faith"], fair_groq["faith"])
make_grouped_4bar("rpt_sidebyside_recall.png", "Context Recall: Original vs Fair Comparison",
    QUESTIONS_SHORT, orig_ollama["recall"], orig_groq["recall"], fair_ollama["recall"], fair_groq["recall"])

print("All charts generated.")

# =========================================================
# BUILD WORD DOCUMENT
# =========================================================
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.bold = True; run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]; cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs: run.font.size = Pt(10)
    doc.add_paragraph()

def add_chart(path):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# =========================================================
# TITLE PAGE
# =========================================================
doc.add_paragraph()
title = doc.add_heading("RAG Evaluation Comparison Report", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.size = Pt(26)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Ollama (llama3.2 — 3B, Local) vs Groq (llama-3.3-70b — API)")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run("Includes original evaluation and fair comparison with same judge model")
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run2.italic = True

doc.add_page_break()

# =========================================================
# SECTION 1: OVERVIEW
# =========================================================
add_heading("1. Overview", level=1)

doc.add_paragraph(
    "This report compares the evaluation results of the same RAG (Retrieval-Augmented Generation) "
    "pipeline using two different LLM backends. Both runs used the same retrieval pipeline "
    "(ChromaDB + sentence-transformers embeddings + cross-encoder reranker), the same 5 questions, "
    "and the same reference documents from HSE medical conditions."
)

add_table(
    ["", "Ollama (Local)", "Groq (API)"],
    [
        ["Model", "llama3.2 (3B params)", "llama-3.3-70b-versatile (70B params)"],
        ["Hosting", "Local (on-device)", "Cloud API (free tier)"],
        ["Role", "RAG generation + evaluation judge", "RAG generation + evaluation judge"],
        ["Questions", "5", "5"],
    ],
)

doc.add_paragraph(
    "The evaluation uses three metrics: Faithfulness (is the answer grounded in context?), "
    "Context Recall (does the context cover the reference answer?), and BLEU Score "
    "(algorithmic word overlap between answer and reference)."
)

# =========================================================
# SECTION 2: ORIGINAL COMPARISON
# =========================================================
doc.add_page_break()
add_heading("2. Original Evaluation Results", level=1)

doc.add_paragraph(
    "In the original evaluation, each model served as both the answer generator and the evaluation judge. "
    "Ollama's llama3.2 (3B) judged its own answers, and Groq's llama-3.3-70b (70B) judged its own answers."
)

add_heading("2.1 Overall Averages", level=2)

add_table(
    ["Metric", "Ollama (llama3.2)", "Groq (llama-3.3-70b)", "Difference"],
    [
        ["Faithfulness", "56.6%", "100.0%", "+43.4%"],
        ["Context Recall", "38.6%", "100.0%", "+61.4%"],
        ["BLEU Score", "51.9%", "63.7%", "+11.8%"],
    ],
)

add_chart("rpt_orig_overall.png")

add_heading("2.2 Faithfulness per Question", level=2)

add_table(
    ["Question", "Ollama", "Groq", "Diff"],
    [[QUESTIONS_FULL[i], f"{orig_ollama['faith'][i]*100:.1f}%", f"{orig_groq['faith'][i]*100:.1f}%",
      f"{(orig_groq['faith'][i]-orig_ollama['faith'][i])*100:+.1f}%"] for i in range(5)],
)
add_chart("rpt_orig_faith.png")

add_heading("2.3 Context Recall per Question", level=2)

add_table(
    ["Question", "Ollama", "Groq", "Diff"],
    [[QUESTIONS_FULL[i], f"{orig_ollama['recall'][i]*100:.1f}%", f"{orig_groq['recall'][i]*100:.1f}%",
      f"{(orig_groq['recall'][i]-orig_ollama['recall'][i])*100:+.1f}%"] for i in range(5)],
)
add_chart("rpt_orig_recall.png")

add_heading("2.4 BLEU Score per Question", level=2)

add_table(
    ["Question", "Ollama", "Groq", "Diff"],
    [[QUESTIONS_FULL[i], f"{orig_ollama['bleu'][i]*100:.1f}%", f"{orig_groq['bleu'][i]*100:.1f}%",
      f"{(orig_groq['bleu'][i]-orig_ollama['bleu'][i])*100:+.1f}%"] for i in range(5)],
)
add_chart("rpt_orig_bleu.png")

# =========================================================
# SECTION 3: FAIR COMPARISON
# =========================================================
doc.add_page_break()
add_heading("3. Fair Comparison (Same Judge: Groq 70B)", level=1)

doc.add_paragraph(
    "To isolate answer quality from judge quality, we re-evaluated both sets of answers "
    "using the same judge model: Groq's llama-3.3-70b-versatile (70B parameters). "
    "The original generated answers from both Ollama and Groq were kept unchanged — "
    "only the evaluation judge was standardised."
)

add_heading("3.1 Overall Averages (Same Judge)", level=2)

add_table(
    ["Metric", "Ollama Answers", "Groq Answers", "Difference"],
    [
        ["Faithfulness", f"{avg(fair_ollama['faith'])*100:.1f}%", f"{avg(fair_groq['faith'])*100:.1f}%",
         f"{(avg(fair_groq['faith'])-avg(fair_ollama['faith']))*100:+.1f}%"],
        ["Context Recall", f"{avg(fair_ollama['recall'])*100:.1f}%", f"{avg(fair_groq['recall'])*100:.1f}%",
         f"{(avg(fair_groq['recall'])-avg(fair_ollama['recall']))*100:+.1f}%"],
        ["BLEU Score", f"{avg(fair_ollama['bleu'])*100:.1f}%", f"{avg(fair_groq['bleu'])*100:.1f}%",
         f"{(avg(fair_groq['bleu'])-avg(fair_ollama['bleu']))*100:+.1f}%"],
    ],
)

add_chart("rpt_fair_overall.png")

add_heading("3.2 Faithfulness per Question (Same Judge)", level=2)

add_table(
    ["Question", "Ollama", "Groq", "Diff"],
    [[QUESTIONS_FULL[i], f"{fair_ollama['faith'][i]*100:.1f}%", f"{fair_groq['faith'][i]*100:.1f}%",
      f"{(fair_groq['faith'][i]-fair_ollama['faith'][i])*100:+.1f}%"] for i in range(5)],
)
add_chart("rpt_fair_faith.png")

add_heading("3.3 Context Recall per Question (Same Judge)", level=2)

add_table(
    ["Question", "Ollama", "Groq", "Diff"],
    [[QUESTIONS_FULL[i], f"{fair_ollama['recall'][i]*100:.1f}%", f"{fair_groq['recall'][i]*100:.1f}%",
      f"{(fair_groq['recall'][i]-fair_ollama['recall'][i])*100:+.1f}%"] for i in range(5)],
)
add_chart("rpt_fair_recall.png")

add_heading("3.4 BLEU Score per Question (Same Judge)", level=2)

add_table(
    ["Question", "Ollama", "Groq", "Diff"],
    [[QUESTIONS_FULL[i], f"{fair_ollama['bleu'][i]*100:.1f}%", f"{fair_groq['bleu'][i]*100:.1f}%",
      f"{(fair_groq['bleu'][i]-fair_ollama['bleu'][i])*100:+.1f}%"] for i in range(5)],
)
add_chart("rpt_fair_bleu.png")

# =========================================================
# SECTION 4: SIDE-BY-SIDE ORIGINAL VS FAIR
# =========================================================
doc.add_page_break()
add_heading("4. Original vs Fair: Side-by-Side", level=1)

doc.add_paragraph(
    "These charts show all four score variants per question — original self-judged scores "
    "alongside the fair Groq-judged scores — making the impact of judge quality immediately visible."
)

add_heading("4.1 Faithfulness: Original vs Fair", level=2)
add_chart("rpt_sidebyside_faith.png")

add_heading("4.2 Context Recall: Original vs Fair", level=2)
add_chart("rpt_sidebyside_recall.png")

# Impact table
add_heading("4.3 Impact of Judge Change on Ollama Scores", level=2)

doc.add_paragraph(
    "This table shows how Ollama's scores changed when re-evaluated by the Groq 70B judge "
    "instead of its own 3B judge."
)

add_table(
    ["Metric", "Ollama (self-judged)", "Ollama (Groq-judged)", "Change"],
    [
        ["Faithfulness", "56.6%", f"{avg(fair_ollama['faith'])*100:.1f}%",
         f"{avg(fair_ollama['faith'])*100-56.6:+.1f}%"],
        ["Context Recall", "38.6%", f"{avg(fair_ollama['recall'])*100:.1f}%",
         f"{avg(fair_ollama['recall'])*100-38.6:+.1f}%"],
        ["BLEU Score", "51.9%", f"{avg(fair_ollama['bleu'])*100:.1f}%",
         f"{avg(fair_ollama['bleu'])*100-51.9:+.1f}%"],
    ],
)

# =========================================================
# SECTION 5: ANALYSIS
# =========================================================
doc.add_page_break()
add_heading("5. Analysis: Why the Results Differ", level=1)

add_heading("5.1 Judge Model Quality — The Main Factor", level=2)
doc.add_paragraph(
    "Faithfulness and context recall are scored by an LLM acting as a judge. In the original evaluation, "
    "Ollama used its own llama3.2 (3B parameters) as the judge, while Groq used llama-3.3-70b (70B parameters)."
)
doc.add_paragraph(
    "The 3B model is a weak evaluator. It struggles to properly assess whether a response is grounded "
    "in the provided context, often giving inconsistent and lower scores even when the answer is clearly "
    "faithful. The 70B model understands the evaluation task much better and scores more accurately."
)
doc.add_paragraph(
    "Evidence: When we re-judged Ollama's answers with the Groq 70B judge, faithfulness jumped from "
    f"56.6% to {avg(fair_ollama['faith'])*100:.1f}% and context recall from 38.6% to "
    f"{avg(fair_ollama['recall'])*100:.1f}%. The answers didn't change — only the judge got smarter."
)

add_heading("5.2 Generation Quality — A Secondary Factor", level=2)
doc.add_paragraph(
    "The 70B model also generates slightly better RAG answers than the 3B model. Examples:"
)
doc.add_paragraph(
    'Melanoma: Groq\'s answer included "Melanoma is not always preventable, but..." which closely '
    'matches the reference. Ollama\'s version said "Melanoma can be prevented by..." — a subtle '
    "but meaningful inaccuracy that omits the nuance from the source material.",
    style="List Bullet",
)
doc.add_paragraph(
    "MRSA: Groq's answer stayed focused on \"how you get MRSA\" (the actual question), while "
    "Ollama's answer drifted into treatment details from other context chunks, reducing faithfulness.",
    style="List Bullet",
)
doc.add_paragraph(
    "Type 1 Diabetes: Both models acknowledged the thin context, but Groq's response was more "
    "concise and cited the source, while Ollama added unsupported advice about \"exploring additional resources.\"",
    style="List Bullet",
)

add_heading("5.3 BLEU Tells the Objective Story", level=2)
doc.add_paragraph(
    "BLEU is computed algorithmically with no LLM involvement, making it the most objective metric. "
    "The BLEU improvement was moderate (51.9% → 63.7%), confirming that the 70B model generates "
    "better answers, but not as dramatically as the original LLM-judged metrics suggested."
)
doc.add_paragraph(
    "BLEU scores are identical in both the original and fair comparison because they don't depend "
    "on the judge model at all."
)

add_heading("5.4 Type 1 Diabetes — Low for Both", level=2)
doc.add_paragraph(
    'Both models scored low on BLEU for this question because the reference answer is extremely '
    'short: "Advice on avoiding complications of type 1 diabetes." There simply isn\'t enough '
    "ground-truth content for either model to match against. This is a data quality issue, "
    "not a model quality issue."
)

# =========================================================
# SECTION 6: KEY TAKEAWAYS
# =========================================================
add_heading("6. Key Takeaways", level=1)

doc.add_paragraph(
    "Using a small model (3B) as both generator and judge produces unreliable evaluation scores. "
    "The 3B judge underscored Ollama's answers by over 30 percentage points on faithfulness.",
    style="List Number",
)
doc.add_paragraph(
    "The 70B model via Groq API provides better generation quality and much more reliable evaluation. "
    "The fair comparison shows the actual generation gap is modest (~12% on faithfulness, ~12% on BLEU).",
    style="List Number",
)
doc.add_paragraph(
    "BLEU scores are the most objective metric since they don't depend on LLM judgment. "
    "They should be used alongside LLM-judged metrics to cross-validate results.",
    style="List Number",
)
doc.add_paragraph(
    "For reliable evaluation, always use a strong model as the judge — ideally separate from "
    "the generation model. This avoids self-evaluation bias.",
    style="List Number",
)
doc.add_paragraph(
    "Both models performed well on context recall when fairly judged (100%), indicating the "
    "retrieval pipeline itself is working effectively.",
    style="List Number",
)

# =========================================================
# SAVE
# =========================================================
doc.save("Full_Evaluation_Comparison_Report.docx")
print("Saved: Full_Evaluation_Comparison_Report.docx")
