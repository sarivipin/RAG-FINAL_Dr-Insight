"""
Fair Comparison: Re-evaluate Ollama-generated answers using Groq (llama-3.3-70b) as judge.
This ensures both Ollama and Groq answers are scored by the same judge model.
"""
import os
import re
import json
import time
from typing import List, Dict
from dotenv import load_dotenv

load_dotenv()

from langchain_groq import ChatGroq

# =========================================================
# CONFIG
# =========================================================
GROQ_MODEL = "llama-3.3-70b-versatile"
OLLAMA_DETAILS = "ragas_details.json"
GROQ_DETAILS = "ragas_details_groq.json"

RESULTS_CSV = "ragas_results_fair_comparison.csv"
DETAILS_JSON = "ragas_details_fair_comparison.json"
REPORT_DOCX = "Fair_Comparison_Report.docx"

MAX_RETRIES = 5
INITIAL_BACKOFF = 15


def _make_judge() -> ChatGroq:
    return ChatGroq(
        model=GROQ_MODEL,
        api_key=os.getenv("GROQ_API_KEY"),
        temperature=0.0,
    )


def invoke_with_retry(llm, prompt: str) -> str:
    for attempt in range(MAX_RETRIES):
        try:
            result = llm.invoke(prompt)
            return result.content
        except Exception as e:
            if "429" in str(e) or "rate" in str(e).lower():
                wait = INITIAL_BACKOFF * (2 ** attempt)
                print(f"    Rate limited. Waiting {wait}s ({attempt + 1}/{MAX_RETRIES})...")
                time.sleep(wait)
            else:
                raise
    raise RuntimeError(f"Failed after {MAX_RETRIES} retries.")


# =========================================================
# METRICS
# =========================================================
def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def compute_bleu(reference: str, candidate: str) -> float:
    ref_tokens = normalize_text(reference).split()
    cand_tokens = normalize_text(candidate).split()
    if not cand_tokens or not ref_tokens:
        return 0.0
    from collections import Counter
    clipped = 0
    ref_counts = Counter(ref_tokens)
    cand_counts = Counter(cand_tokens)
    for token, count in cand_counts.items():
        clipped += min(count, ref_counts.get(token, 0))
    precision = clipped / len(cand_tokens) if cand_tokens else 0
    brevity = min(1.0, len(cand_tokens) / len(ref_tokens)) if ref_tokens else 0
    return precision * brevity


def evaluate_faithfulness(llm, response: str, contexts: List[str]) -> float:
    if not contexts or not response.strip():
        return 0.0
    context_text = "\n\n".join(contexts[:3])
    prompt = f"""You are an evaluation judge. Rate how faithful the following response is to the provided context.
A faithful response only contains information that can be found in or inferred from the context.

Context:
{context_text}

Response:
{response}

Rate faithfulness as a score from 0.0 to 1.0 where:
- 1.0 = completely faithful, all claims supported by context
- 0.0 = completely unfaithful, no claims supported

Reply with ONLY a number between 0.0 and 1.0, nothing else."""
    try:
        content = invoke_with_retry(llm, prompt)
        score = float(re.search(r"([01]\.?\d*)", content).group(1))
        return min(max(score, 0.0), 1.0)
    except Exception:
        return 0.0


def evaluate_context_recall(llm, response: str, contexts: List[str], reference: str) -> float:
    if not contexts or not reference.strip():
        return 0.0
    context_text = "\n\n".join(contexts[:3])
    prompt = f"""You are an evaluation judge. Rate how well the retrieved context covers the information needed to answer correctly.

Reference answer (ground truth):
{reference}

Retrieved context:
{context_text}

Rate context recall as a score from 0.0 to 1.0 where:
- 1.0 = context contains all information from the reference
- 0.0 = context contains none of the reference information

Reply with ONLY a number between 0.0 and 1.0, nothing else."""
    try:
        content = invoke_with_retry(llm, prompt)
        score = float(re.search(r"([01]\.?\d*)", content).group(1))
        return min(max(score, 0.0), 1.0)
    except Exception:
        return 0.0


# =========================================================
# RE-EVALUATE EXISTING RESULTS WITH GROQ JUDGE
# =========================================================
def rejudge(rows: List[Dict], label: str) -> List[Dict]:
    llm = _make_judge()
    scored = []
    for idx, row in enumerate(rows, start=1):
        print(f"  [{label}] Judging [{idx}/{len(rows)}] {row['user_input'][:60]}...")
        f = evaluate_faithfulness(llm, row["response"], row["retrieved_contexts"])
        cr = evaluate_context_recall(llm, row["response"], row["retrieved_contexts"], row["reference"])
        bleu = compute_bleu(row["reference"], row["response"])
        scored.append({
            "user_input": row["user_input"],
            "response": row["response"],
            "faithfulness": f,
            "context_recall": cr,
            "bleu": bleu,
        })
    return scored


# =========================================================
# GENERATE WORD REPORT
# =========================================================
def generate_report(ollama_scored, groq_scored):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # --- Compute averages ---
    avg = lambda lst, key: sum(r[key] for r in lst) / len(lst) if lst else 0

    ollama_avg = [avg(ollama_scored, "faithfulness"), avg(ollama_scored, "context_recall"), avg(ollama_scored, "bleu")]
    groq_avg = [avg(groq_scored, "faithfulness"), avg(groq_scored, "context_recall"), avg(groq_scored, "bleu")]

    questions_short = [
        "Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention",
        "Type 1\nDiabetes", "MRSA",
    ]
    c_ollama = "#4A90D9"
    c_groq = "#E8833A"

    # --- Chart: Overall ---
    fig, ax = plt.subplots(figsize=(8, 5))
    metrics = ["Faithfulness", "Context Recall", "BLEU Score"]
    x = np.arange(len(metrics))
    w = 0.32
    b1 = ax.bar(x - w/2, [v*100 for v in ollama_avg], w, label="Ollama answers", color=c_ollama)
    b2 = ax.bar(x + w/2, [v*100 for v in groq_avg], w, label="Groq answers", color=c_groq)
    ax.set_ylabel("Score (%)")
    ax.set_title("Fair Comparison: Same Judge (Groq 70B) for Both", fontsize=13, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(metrics)
    ax.set_ylim(0, 115); ax.legend(); ax.grid(axis="y", alpha=0.3)
    for b in b1:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=9)
    for b in b2:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=9)
    plt.tight_layout()
    plt.savefig("chart_fair_overall.png", dpi=150); plt.close()

    # --- Chart: Per-question grouped ---
    for metric_key, metric_label in [("faithfulness", "Faithfulness"), ("context_recall", "Context Recall"), ("bleu", "BLEU Score")]:
        fig, ax = plt.subplots(figsize=(10, 5))
        x = np.arange(len(questions_short))
        o_vals = [r[metric_key]*100 for r in ollama_scored]
        g_vals = [r[metric_key]*100 for r in groq_scored]
        ax.bar(x - w/2, o_vals, w, label="Ollama answers", color=c_ollama)
        ax.bar(x + w/2, g_vals, w, label="Groq answers", color=c_groq)
        ax.set_ylabel("Score (%)"); ax.set_title(f"{metric_label} per Question (Same Judge)", fontsize=13, fontweight="bold")
        ax.set_xticks(x); ax.set_xticklabels(questions_short, fontsize=9)
        ax.set_ylim(0, 115); ax.legend(); ax.grid(axis="y", alpha=0.3)
        plt.tight_layout()
        plt.savefig(f"chart_fair_{metric_key}.png", dpi=150); plt.close()

    # --- Word Document ---
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    def add_table(headers, rows):
        table = doc.add_table(rows=1+len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]; cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs: run.bold = True; run.font.size = Pt(10)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx+1].cells[c_idx]; cell.text = str(val)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs: run.font.size = Pt(10)
        doc.add_paragraph()

    # Title
    title = doc.add_heading("Fair Comparison Report", level=0)
    for run in title.runs: run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    sub = doc.add_paragraph("Both answer sets judged by the same model: Groq llama-3.3-70b-versatile")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs: run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # Approach
    add_heading("1. Approach", level=1)
    doc.add_paragraph(
        "The original evaluation used each model as both generator and judge, making comparison unfair. "
        "In this fair comparison, we keep the original generated answers from both Ollama and Groq, "
        "but re-evaluate ALL answers using the same judge: Groq's llama-3.3-70b-versatile (70B parameters). "
        "This isolates the answer quality from the judge quality."
    )

    # Overall
    add_heading("2. Overall Results (Same Judge)", level=1)
    add_table(
        ["Metric", "Ollama Answers", "Groq Answers", "Difference"],
        [
            ["Faithfulness", f"{ollama_avg[0]*100:.1f}%", f"{groq_avg[0]*100:.1f}%", f"{(groq_avg[0]-ollama_avg[0])*100:+.1f}%"],
            ["Context Recall", f"{ollama_avg[1]*100:.1f}%", f"{groq_avg[1]*100:.1f}%", f"{(groq_avg[1]-ollama_avg[1])*100:+.1f}%"],
            ["BLEU Score", f"{ollama_avg[2]*100:.1f}%", f"{groq_avg[2]*100:.1f}%", f"{(groq_avg[2]-ollama_avg[2])*100:+.1f}%"],
        ],
    )
    doc.add_picture("chart_fair_overall.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Per-question tables + charts
    add_heading("3. Per-Question Breakdown", level=1)

    for metric_key, metric_label, chart_file in [
        ("faithfulness", "Faithfulness", "chart_fair_faithfulness.png"),
        ("context_recall", "Context Recall", "chart_fair_context_recall.png"),
        ("bleu", "BLEU Score", "chart_fair_bleu.png"),
    ]:
        add_heading(f"3.{['faithfulness','context_recall','bleu'].index(metric_key)+1} {metric_label}", level=2)
        q_names = [
            "Symptoms of bacterial vaginosis", "Causes of varicose veins",
            "Preventing melanoma", "Living with type 1 diabetes", "MRSA",
        ]
        rows = []
        for i in range(5):
            o = ollama_scored[i][metric_key]*100
            g = groq_scored[i][metric_key]*100
            rows.append([q_names[i], f"{o:.1f}%", f"{g:.1f}%", f"{g-o:+.1f}%"])
        add_table(["Question", "Ollama", "Groq", "Diff"], rows)
        doc.add_picture(chart_file, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Analysis
    add_heading("4. What This Tells Us", level=1)
    doc.add_paragraph(
        "With the same judge model evaluating both sets of answers, we can now see the true "
        "difference in answer quality between the 3B and 70B generation models, free from judge bias."
    )
    doc.add_paragraph(
        "The BLEU scores remain unchanged since they are purely algorithmic. "
        "The faithfulness and context recall scores now reflect actual answer quality differences "
        "rather than judge capability differences.",
    )

    # Comparison with original
    add_heading("5. Original vs Fair Comparison", level=1)
    add_table(
        ["Metric", "Ollama (self-judged)", "Ollama (Groq-judged)", "Change"],
        [
            ["Faithfulness", "56.6%", f"{ollama_avg[0]*100:.1f}%", f"{ollama_avg[0]*100-56.6:+.1f}%"],
            ["Context Recall", "38.6%", f"{ollama_avg[1]*100:.1f}%", f"{ollama_avg[1]*100-38.6:+.1f}%"],
            ["BLEU Score", "51.9%", f"{ollama_avg[2]*100:.1f}%", f"{ollama_avg[2]*100-51.9:+.1f}%"],
        ],
    )
    doc.add_paragraph(
        "This table shows how much the Ollama scores change when judged by a better model. "
        "The difference highlights how unreliable the small 3B model was as an evaluator."
    )

    doc.save(REPORT_DOCX)
    print(f"\nSaved: {REPORT_DOCX}")


# =========================================================
# MAIN
# =========================================================
def main():
    print("Loading Ollama results from ragas_details.json...")
    with open(OLLAMA_DETAILS) as f:
        ollama_rows = json.load(f)

    print("Loading Groq results from ragas_details_groq.json...")
    with open(GROQ_DETAILS) as f:
        groq_rows = json.load(f)

    print(f"\nRe-judging {len(ollama_rows)} Ollama answers with Groq 70B judge...")
    ollama_scored = rejudge(ollama_rows, "Ollama")

    print(f"\nRe-judging {len(groq_rows)} Groq answers with Groq 70B judge...")
    groq_scored = rejudge(groq_rows, "Groq")

    # Save CSV
    import pandas as pd
    all_rows = []
    for o, g in zip(ollama_scored, groq_scored):
        all_rows.append({
            "question": o["user_input"],
            "ollama_faithfulness": o["faithfulness"],
            "groq_faithfulness": g["faithfulness"],
            "ollama_context_recall": o["context_recall"],
            "groq_context_recall": g["context_recall"],
            "ollama_bleu": o["bleu"],
            "groq_bleu": g["bleu"],
        })
    pd.DataFrame(all_rows).to_csv(RESULTS_CSV, index=False)
    print(f"\nSaved: {RESULTS_CSV}")

    # Save JSON
    with open(DETAILS_JSON, "w") as f:
        json.dump({"ollama_scored": ollama_scored, "groq_scored": groq_scored}, f, indent=2)
    print(f"Saved: {DETAILS_JSON}")

    # Print summary
    avg = lambda lst, k: sum(r[k] for r in lst) / len(lst)
    print("\n" + "=" * 60)
    print("FAIR COMPARISON — Same Judge (Groq 70B)")
    print("=" * 60)
    print(f"{'Metric':<20} {'Ollama Answers':>15} {'Groq Answers':>15}")
    print("-" * 50)
    for k in ["faithfulness", "context_recall", "bleu"]:
        print(f"{k:<20} {avg(ollama_scored, k)*100:>14.1f}% {avg(groq_scored, k)*100:>14.1f}%")
    print("=" * 60)

    # Generate Word report
    print("\nGenerating Word report with charts...")
    generate_report(ollama_scored, groq_scored)


if __name__ == "__main__":
    main()
