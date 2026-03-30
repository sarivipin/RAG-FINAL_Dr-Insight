from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- Styles ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# =========================================================
# TITLE
# =========================================================
title = doc.add_heading("RAG Evaluation Comparison Report", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph("Ollama (llama3.2 - 3B, Local) vs Groq (llama-3.3-70b, API)")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# =========================================================
# SECTION 1: Overview
# =========================================================
add_heading("1. Overview", level=1)

doc.add_paragraph(
    "This report compares the evaluation results of the same RAG (Retrieval-Augmented Generation) "
    "pipeline using two different LLM backends. Both runs used the same retrieval pipeline "
    "(ChromaDB + sentence-transformers + cross-encoder reranker), same 5 questions, and same reference documents."
)

add_table(
    ["", "Ollama", "Groq"],
    [
        ["Model", "llama3.2 (3B params)", "llama-3.3-70b-versatile (70B params)"],
        ["Hosting", "Local (on-device)", "Cloud API (free tier)"],
        ["Used for", "RAG generation + evaluation", "RAG generation + evaluation"],
        ["Questions evaluated", "5", "5"],
    ],
)

# =========================================================
# SECTION 2: Overall Results
# =========================================================
add_heading("2. Overall Results", level=1)

add_table(
    ["Metric", "Ollama (llama3.2)", "Groq (llama-3.3-70b)", "Difference"],
    [
        ["Faithfulness", "56.6%", "100.0%", "+43.4%"],
        ["Context Recall", "38.6%", "100.0%", "+61.4%"],
        ["BLEU Score", "51.9%", "63.7%", "+11.8%"],
    ],
)

doc.add_picture("chart_overall_comparison.png", width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# =========================================================
# SECTION 3: Per-Question Breakdown
# =========================================================
add_heading("3. Per-Question Breakdown", level=1)

# Faithfulness
add_heading("3.1 Faithfulness", level=2)
doc.add_paragraph(
    "Faithfulness measures whether the generated answer only contains information "
    "that is supported by the retrieved context."
)

add_table(
    ["Question", "Ollama", "Groq"],
    [
        ["Symptoms of bacterial vaginosis", "80.0%", "100.0%"],
        ["Causes of varicose veins", "80.0%", "100.0%"],
        ["Preventing melanoma", "33.0%", "100.0%"],
        ["Living with type 1 diabetes", "50.0%", "100.0%"],
        ["MRSA", "40.0%", "100.0%"],
    ],
)

doc.add_picture("chart_faithfulness_per_question.png", width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Context Recall
add_heading("3.2 Context Recall", level=2)
doc.add_paragraph(
    "Context recall measures how well the retrieved context covers the information "
    "in the ground-truth reference answer."
)

add_table(
    ["Question", "Ollama", "Groq"],
    [
        ["Symptoms of bacterial vaginosis", "40.0%", "100.0%"],
        ["Causes of varicose veins", "40.0%", "100.0%"],
        ["Preventing melanoma", "40.0%", "100.0%"],
        ["Living with type 1 diabetes", "33.0%", "100.0%"],
        ["MRSA", "40.0%", "100.0%"],
    ],
)

doc.add_picture("chart_context_recall_per_question.png", width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# BLEU
add_heading("3.3 BLEU Score", level=2)
doc.add_paragraph(
    "BLEU is a purely algorithmic metric (no LLM involved) that measures word overlap "
    "between the generated answer and the reference answer."
)

add_table(
    ["Question", "Ollama", "Groq"],
    [
        ["Symptoms of bacterial vaginosis", "91.1%", "86.8%"],
        ["Causes of varicose veins", "45.9%", "36.1%"],
        ["Preventing melanoma", "78.5%", "96.9%"],
        ["Living with type 1 diabetes", "9.6%", "20.0%"],
        ["MRSA", "34.2%", "78.6%"],
    ],
)

doc.add_picture("chart_bleu_per_question.png", width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# =========================================================
# SECTION 4: Analysis
# =========================================================
add_heading("4. Why Such a Big Difference?", level=1)

add_heading("4.1 Judge Model Quality (Main Factor)", level=2)
doc.add_paragraph(
    "Faithfulness and context recall are scored by the LLM acting as a judge. "
    "The 3B model (llama3.2) is a weak evaluator — it inconsistently assesses whether "
    "responses are grounded in context, often giving lower scores even when the answer "
    "is clearly faithful. The 70B model (llama-3.3-70b) understands the evaluation task "
    "much better and scores more accurately."
)
doc.add_paragraph(
    "This is the primary reason for the large gap in faithfulness (56.6% → 100%) "
    "and context recall (38.6% → 100%). The judge got smarter, not just the answers."
)

add_heading("4.2 Generation Quality (Secondary Factor)", level=2)
doc.add_paragraph(
    "The 70B model also generates slightly better RAG answers. For example:"
)
doc.add_paragraph(
    "Melanoma: Groq included \"Melanoma is not always preventable, but...\" "
    "(matching the reference closely), while Ollama said \"Melanoma can be prevented by...\" "
    "— a subtle but meaningful inaccuracy.",
    style="List Bullet",
)
doc.add_paragraph(
    "MRSA: Groq's answer was more concise and closer to the source material, "
    "staying focused on the \"how you get MRSA\" topic without drifting into treatment details.",
    style="List Bullet",
)

add_heading("4.3 BLEU Tells the Real Story", level=2)
doc.add_paragraph(
    "Since BLEU is computed algorithmically with no LLM bias, it shows the actual answer "
    "quality difference more honestly. The improvement is moderate (51.9% → 63.7%), "
    "confirming that the answer quality improved but not as dramatically as the "
    "LLM-judged metrics suggest."
)

add_heading("4.4 Type 1 Diabetes — Low for Both", level=2)
doc.add_paragraph(
    "Both models scored low on BLEU for this question because the reference answer is "
    "extremely short (\"Advice on avoiding complications of type 1 diabetes\"). There simply "
    "isn't enough ground-truth content for either model to match against."
)

# =========================================================
# SECTION 5: Key Takeaways
# =========================================================
add_heading("5. Key Takeaways", level=1)

doc.add_paragraph(
    "Using a small model (3B) as both generator and judge produces unreliable evaluation scores.",
    style="List Number",
)
doc.add_paragraph(
    "The 70B model via Groq API provides significantly better generation and more reliable evaluation.",
    style="List Number",
)
doc.add_paragraph(
    "BLEU scores are the most objective metric in this comparison since they don't depend on LLM judgment.",
    style="List Number",
)
doc.add_paragraph(
    "For a fair comparison, the same judge model should evaluate both sets of answers "
    "(e.g., use Groq's 70B to judge Ollama's answers too).",
    style="List Number",
)

# =========================================================
# SAVE
# =========================================================
doc.save("Evaluation_Comparison_Report.docx")
print("Saved: Evaluation_Comparison_Report.docx")
