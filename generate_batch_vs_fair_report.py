"""
Compare batch_evaluation.py (token-overlap metrics, Ollama generation)
vs fair comparison (LLM-judged metrics, Groq 70B judge).
Both use Ollama-generated answers for the common metric: faithfulness.
"""
import json
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

C_BATCH = "#2ECC71"
C_FAIR_OLLAMA = "#4A90D9"
C_FAIR_GROQ = "#E8833A"

QUESTIONS_SHORT = ["Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention", "Type 1\nDiabetes", "MRSA"]
QUESTIONS_FULL = [
    "Symptoms of bacterial vaginosis",
    "Causes of varicose veins",
    "Preventing melanoma",
    "Living with type 1 diabetes",
    "MRSA",
]

# ---- Batch evaluation scores (token-overlap, Ollama-generated) ----
batch_faith = [0.95, 0.935, 0.96, 0.319, 0.667]
batch_coverage = [0.266, 0.218, 0.289, 0.344, 0.352]
batch_relevance = [0.571, 0.250, 0.714, 0.800, 0.545]

# ---- Fair comparison scores (LLM-judged by Groq 70B) ----
# Ollama answers judged by Groq
fair_ollama_faith = [1.0, 1.0, 1.0, 0.8, 0.6]
fair_ollama_recall = [1.0, 1.0, 1.0, 1.0, 1.0]
fair_ollama_bleu = [0.911, 0.459, 0.785, 0.096, 0.342]

# Groq answers judged by Groq
fair_groq_faith = [1.0, 1.0, 1.0, 1.0, 1.0]
fair_groq_recall = [1.0, 1.0, 1.0, 1.0, 1.0]
fair_groq_bleu = [0.868, 0.361, 0.969, 0.200, 0.786]

avg = lambda lst: sum(lst) / len(lst)

# =========================================================
# GENERATE CHARTS
# =========================================================

# Chart 1: Faithfulness comparison — batch (token-overlap) vs fair (LLM-judged) for Ollama answers
def chart_faith_comparison():
    fig, ax = plt.subplots(figsize=(10, 5))
    x = np.arange(len(QUESTIONS_SHORT))
    w = 0.28
    ax.bar(x - w, [v*100 for v in batch_faith], w, label="Batch (token-overlap)", color=C_BATCH)
    ax.bar(x, [v*100 for v in fair_ollama_faith], w, label="Fair: Ollama (Groq-judged)", color=C_FAIR_OLLAMA)
    ax.bar(x + w, [v*100 for v in fair_groq_faith], w, label="Fair: Groq (Groq-judged)", color=C_FAIR_GROQ)
    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title("Faithfulness: Batch (Token-Overlap) vs Fair Comparison (LLM-Judged)", fontsize=12, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(QUESTIONS_SHORT, fontsize=9)
    ax.set_ylim(0, 118); ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig("bvf_faithfulness.png", dpi=150); plt.close()

# Chart 2: Overall averages — all metrics side by side
def chart_overall():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    metrics = ["Faithfulness", "Coverage /\nContext Recall", "Relevance /\nBLEU Score"]
    x = np.arange(len(metrics))
    w = 0.22

    batch_vals = [avg(batch_faith)*100, avg(batch_coverage)*100, avg(batch_relevance)*100]
    fair_o_vals = [avg(fair_ollama_faith)*100, avg(fair_ollama_recall)*100, avg(fair_ollama_bleu)*100]
    fair_g_vals = [avg(fair_groq_faith)*100, avg(fair_groq_recall)*100, avg(fair_groq_bleu)*100]

    b1 = ax.bar(x - w, batch_vals, w, label="Batch (token-overlap, Ollama)", color=C_BATCH)
    b2 = ax.bar(x, fair_o_vals, w, label="Fair: Ollama answers (Groq-judged)", color=C_FAIR_OLLAMA)
    b3 = ax.bar(x + w, fair_g_vals, w, label="Fair: Groq answers (Groq-judged)", color=C_FAIR_GROQ)

    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title("Overall Averages: Batch vs Fair Comparison", fontsize=13, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(metrics, fontsize=10)
    ax.set_ylim(0, 118); ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
    for bars in [b1, b2, b3]:
        for b in bars:
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=8)
    plt.tight_layout()
    plt.savefig("bvf_overall.png", dpi=150); plt.close()

# Chart 3: Batch-only metrics (coverage + relevance)
def chart_batch_metrics():
    fig, ax = plt.subplots(figsize=(10, 5))
    x = np.arange(len(QUESTIONS_SHORT))
    w = 0.25
    ax.bar(x - w, [v*100 for v in batch_faith], w, label="Faithfulness", color=C_BATCH)
    ax.bar(x, [v*100 for v in batch_coverage], w, label="Coverage", color="#3498DB")
    ax.bar(x + w, [v*100 for v in batch_relevance], w, label="Relevance", color="#9B59B6")
    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title("Batch Evaluation: All Metrics per Question", fontsize=13, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(QUESTIONS_SHORT, fontsize=9)
    ax.set_ylim(0, 118); ax.legend(fontsize=10); ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig("bvf_batch_all.png", dpi=150); plt.close()

# Chart 4: Fair comparison all metrics
def chart_fair_metrics():
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    x = np.arange(len(QUESTIONS_SHORT))
    w = 0.32

    # Ollama answers
    ax = axes[0]
    ax.bar(x - w, [v*100 for v in fair_ollama_faith], w, label="Faithfulness", color=C_FAIR_OLLAMA)
    ax.bar(x, [v*100 for v in fair_ollama_recall], w, label="Context Recall", color="#5DADE2")
    ax.bar(x + w, [v*100 for v in fair_ollama_bleu], w, label="BLEU", color="#85C1E9")
    ax.set_title("Fair: Ollama Answers (Groq-judged)", fontsize=11, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(QUESTIONS_SHORT, fontsize=8)
    ax.set_ylim(0, 118); ax.legend(fontsize=8); ax.grid(axis="y", alpha=0.3)

    # Groq answers
    ax = axes[1]
    ax.bar(x - w, [v*100 for v in fair_groq_faith], w, label="Faithfulness", color=C_FAIR_GROQ)
    ax.bar(x, [v*100 for v in fair_groq_recall], w, label="Context Recall", color="#F0B27A")
    ax.bar(x + w, [v*100 for v in fair_groq_bleu], w, label="BLEU", color="#F5CBA7")
    ax.set_title("Fair: Groq Answers (Groq-judged)", fontsize=11, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(QUESTIONS_SHORT, fontsize=8)
    ax.set_ylim(0, 118); ax.legend(fontsize=8); ax.grid(axis="y", alpha=0.3)

    plt.tight_layout()
    plt.savefig("bvf_fair_all.png", dpi=150); plt.close()

chart_faith_comparison()
chart_overall()
chart_batch_metrics()
chart_fair_metrics()
print("All charts generated.")

# =========================================================
# BUILD WORD DOCUMENT
# =========================================================
doc = Document()
style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.bold = True; run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]; cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs: run.font.size = Pt(10)
    doc.add_paragraph()

def add_chart(path):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ---- Title ----
doc.add_paragraph()
title = doc.add_heading("Batch Evaluation vs Fair Comparison Report", level=0)
for run in title.runs: run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Token-Overlap Metrics vs LLM-Judged Metrics (Groq 70B)")
run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_page_break()

# ---- Section 1: Overview ----
add_heading("1. Overview", level=1)

doc.add_paragraph(
    "This report compares two different evaluation approaches applied to the same RAG pipeline:"
)

add_table(
    ["", "Batch Evaluation", "Fair Comparison"],
    [
        ["File", "batch_evaluation.py", "evaluation_fair_comparison.py"],
        ["Evaluation Method", "Token-overlap (algorithmic)", "LLM-as-judge (Groq 70B)"],
        ["Answer Generator", "Ollama (llama3.2, 3B)", "Ollama + Groq (both evaluated)"],
        ["Faithfulness", "Token overlap between answer & context", "LLM judges if answer is grounded in context"],
        ["Coverage / Recall", "Token overlap (context → answer)", "LLM judges if context covers reference"],
        ["Relevance / BLEU", "Query-answer token overlap", "Word overlap with reference answer"],
    ],
)

doc.add_paragraph(
    "The batch evaluation uses simple token-overlap metrics — fast and deterministic but shallow. "
    "The fair comparison uses an LLM (Groq 70B) as a judge — slower but captures semantic meaning."
)

# ---- Section 2: Batch Evaluation Results ----
doc.add_page_break()
add_heading("2. Batch Evaluation Results (Token-Overlap)", level=1)

doc.add_paragraph(
    "The batch evaluation uses Ollama (llama3.2) to generate answers and scores them using "
    "simple token-overlap metrics. No LLM is involved in scoring."
)

add_heading("2.1 Overall Averages", level=2)
add_table(
    ["Metric", "Average Score"],
    [
        ["Faithfulness (answer ∩ context / answer)", f"{avg(batch_faith)*100:.1f}%"],
        ["Coverage (answer ∩ context / context)", f"{avg(batch_coverage)*100:.1f}%"],
        ["Relevance (query ∩ answer / query)", f"{avg(batch_relevance)*100:.1f}%"],
    ],
)

add_heading("2.2 Per-Question Scores", level=2)
add_table(
    ["Question", "Faithfulness", "Coverage", "Relevance"],
    [[QUESTIONS_FULL[i], f"{batch_faith[i]*100:.1f}%", f"{batch_coverage[i]*100:.1f}%",
      f"{batch_relevance[i]*100:.1f}%"] for i in range(5)],
)
add_chart("bvf_batch_all.png")

# ---- Section 3: Fair Comparison Results ----
doc.add_page_break()
add_heading("3. Fair Comparison Results (LLM-Judged)", level=1)

doc.add_paragraph(
    "The fair comparison re-evaluates both Ollama and Groq generated answers using the same "
    "Groq 70B judge. This provides semantic evaluation rather than surface-level token matching."
)

add_heading("3.1 Overall Averages", level=2)
add_table(
    ["Metric", "Ollama Answers", "Groq Answers"],
    [
        ["Faithfulness", f"{avg(fair_ollama_faith)*100:.1f}%", f"{avg(fair_groq_faith)*100:.1f}%"],
        ["Context Recall", f"{avg(fair_ollama_recall)*100:.1f}%", f"{avg(fair_groq_recall)*100:.1f}%"],
        ["BLEU Score", f"{avg(fair_ollama_bleu)*100:.1f}%", f"{avg(fair_groq_bleu)*100:.1f}%"],
    ],
)

add_heading("3.2 Per-Question Scores", level=2)
add_table(
    ["Question", "Ollama Faith.", "Groq Faith.", "Ollama BLEU", "Groq BLEU"],
    [[QUESTIONS_FULL[i],
      f"{fair_ollama_faith[i]*100:.1f}%", f"{fair_groq_faith[i]*100:.1f}%",
      f"{fair_ollama_bleu[i]*100:.1f}%", f"{fair_groq_bleu[i]*100:.1f}%"] for i in range(5)],
)
add_chart("bvf_fair_all.png")

# ---- Section 4: Direct Comparison ----
doc.add_page_break()
add_heading("4. Direct Comparison: Faithfulness", level=1)

doc.add_paragraph(
    "Faithfulness is the one metric both approaches share (though measured differently). "
    "This section compares them directly."
)

add_heading("4.1 Faithfulness: All Three Variants", level=2)
add_table(
    ["Question", "Batch\n(token-overlap)", "Fair: Ollama\n(Groq-judged)", "Fair: Groq\n(Groq-judged)"],
    [[QUESTIONS_FULL[i], f"{batch_faith[i]*100:.1f}%", f"{fair_ollama_faith[i]*100:.1f}%",
      f"{fair_groq_faith[i]*100:.1f}%"] for i in range(5)],
)
add_chart("bvf_faithfulness.png")

add_heading("4.2 Overall Comparison: All Metrics", level=2)
add_table(
    ["Metric", "Batch (Ollama)", "Fair: Ollama\n(Groq-judged)", "Fair: Groq\n(Groq-judged)"],
    [
        ["Faithfulness", f"{avg(batch_faith)*100:.1f}%", f"{avg(fair_ollama_faith)*100:.1f}%", f"{avg(fair_groq_faith)*100:.1f}%"],
        ["Coverage / Context Recall", f"{avg(batch_coverage)*100:.1f}%", f"{avg(fair_ollama_recall)*100:.1f}%", f"{avg(fair_groq_recall)*100:.1f}%"],
        ["Relevance / BLEU", f"{avg(batch_relevance)*100:.1f}%", f"{avg(fair_ollama_bleu)*100:.1f}%", f"{avg(fair_groq_bleu)*100:.1f}%"],
    ],
)
add_chart("bvf_overall.png")

# ---- Section 5: Analysis ----
doc.add_page_break()
add_heading("5. Analysis: Why the Approaches Differ", level=1)

add_heading("5.1 Token-Overlap vs Semantic Understanding", level=2)
doc.add_paragraph(
    "The batch evaluation measures faithfulness by counting how many answer tokens appear in the context. "
    "This is fast and deterministic, but it misses semantic meaning. For example, if the answer says "
    "\"varicose veins are caused by valve problems\" and the context says \"valves stop working properly\", "
    "token overlap may score this lower than an LLM judge that understands they mean the same thing."
)

add_heading("5.2 Coverage vs Context Recall", level=2)
doc.add_paragraph(
    "Batch coverage (29.4%) is much lower than fair context recall (100%). This is because coverage "
    "measures what fraction of context tokens appear in the answer — but the context is much longer "
    "than the answer, so most context tokens naturally won't appear. The LLM judge evaluates whether "
    "the context semantically covers the reference answer, which is a more meaningful question."
)

add_heading("5.3 Relevance vs BLEU", level=2)
doc.add_paragraph(
    "Batch relevance (57.6%) measures query-answer token overlap — does the answer contain the "
    "question's words? BLEU (51.9% Ollama, 63.7% Groq) measures answer-reference word overlap. "
    "These measure different things: relevance checks if the answer addresses the question, "
    "while BLEU checks if the answer matches the ground truth."
)

add_heading("5.4 Type 1 Diabetes — Consistently Weak", level=2)
doc.add_paragraph(
    "Both approaches flagged this question as problematic. Batch faithfulness was 31.9% "
    "(the answer contained many words not in the context), and fair faithfulness was 80% "
    "(the LLM judge noted the answer added unsupported advice). The underlying issue is "
    "thin reference data for this topic."
)

add_heading("5.5 MRSA — Divergent Scores", level=2)
doc.add_paragraph(
    "Batch faithfulness for MRSA was 66.7% while fair comparison gave 60%. Both flagged that "
    "the Ollama answer drifted into treatment details beyond the \"how you get MRSA\" question. "
    "The token-overlap approach caught this because treatment-related tokens weren't in the "
    "primary context chunk."
)

# ---- Section 6: Key Takeaways ----
add_heading("6. Key Takeaways", level=1)

doc.add_paragraph(
    "Token-overlap metrics (batch) are fast, deterministic, and free — but they only capture "
    "surface-level word matching. They work well as a quick sanity check.",
    style="List Number",
)
doc.add_paragraph(
    "LLM-judged metrics (fair comparison) capture semantic meaning and nuance, but depend on "
    "the quality of the judge model and cost API calls.",
    style="List Number",
)
doc.add_paragraph(
    "Both approaches agreed on the weak spots: Type 1 Diabetes (thin data) and MRSA (answer drift). "
    "When both methods flag the same issue, it's a reliable signal.",
    style="List Number",
)
doc.add_paragraph(
    "The ideal evaluation combines both: token-overlap for fast iteration during development, "
    "and LLM-judged metrics for thorough evaluation before deployment.",
    style="List Number",
)
doc.add_paragraph(
    "Coverage (29.4%) should not be compared directly to context recall (100%) — they measure "
    "fundamentally different things despite sounding similar.",
    style="List Number",
)

# ---- Save ----
doc.save("Batch_vs_Fair_Comparison_Report.docx")
print("Saved: Batch_vs_Fair_Comparison_Report.docx")
