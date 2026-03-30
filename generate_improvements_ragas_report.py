"""Generate RAGAS improvements comparison chart and Word report."""
import json
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

with open("improvements_ragas_results.json") as f:
    data = json.load(f)

KEYS = ["baseline", "faiss", "mpnet", "hybrid"]
LABELS = [data[k]["label"] for k in KEYS]
LABELS_SHORT = ["Baseline", "FAISS", "MPNet", "Hybrid"]
COLORS = ["#4A90D9", "#2ECC71", "#E8833A", "#9B59B6"]
QS = ["Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention", "Type 1\nDiabetes", "MRSA"]
QF = [data[KEYS[0]]["per_question"][i]["question"] for i in range(5)]

avgs = {k: [data[k]["avg_faithfulness"], data[k]["avg_context_recall"],
            data[k]["avg_context_precision"], data[k]["avg_retrieve_ms"]] for k in KEYS}
pq = {k: {m: [r[m] for r in data[k]["per_question"]]
           for m in ["faithfulness","context_recall","context_precision"]} for k in KEYS}

# =========================================================
# CHARTS
# =========================================================
# Overall
fig, ax = plt.subplots(figsize=(10, 5.5))
metrics = ["Faithfulness", "Context Recall", "Context Precision"]
x = np.arange(len(metrics)); w = 0.18
for i, (k, l, c) in enumerate(zip(KEYS, LABELS_SHORT, COLORS)):
    vals = [avgs[k][j]*100 for j in range(3)]
    bars = ax.bar(x + (i-1.5)*w, vals, w, label=l, color=c)
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1, f"{b.get_height():.0f}%", ha="center", fontsize=7)
ax.set_ylabel("Score (%)"); ax.set_title("RAG Improvements: RAGAS Evaluation Overall", fontsize=13, fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(metrics); ax.set_ylim(0, 118)
ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("imp_ragas_overall.png", dpi=150); plt.close()

# Speed
fig, ax = plt.subplots(figsize=(8, 5))
vals = [avgs[k][3] for k in KEYS]
bars = ax.bar(np.arange(4), vals, color=COLORS, width=0.5)
for b in bars:
    ax.text(b.get_x()+b.get_width()/2, b.get_height()+20, f"{b.get_height():.0f}ms", ha="center", fontsize=9)
ax.set_ylabel("Time (ms)"); ax.set_title("Average Retrieval Time per Query", fontsize=12, fontweight="bold")
ax.set_xticks(np.arange(4)); ax.set_xticklabels(LABELS_SHORT); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("imp_ragas_speed.png", dpi=150); plt.close()

# Per-question
for metric, label in [("faithfulness","Faithfulness"),("context_recall","Context Recall"),("context_precision","Context Precision")]:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    x = np.arange(5); w = 0.18
    for i, (k, l, c) in enumerate(zip(KEYS, LABELS_SHORT, COLORS)):
        ax.bar(x + (i-1.5)*w, [v*100 for v in pq[k][metric]], w, label=l, color=c)
    ax.set_ylabel("Score (%)"); ax.set_title(f"RAGAS {label} per Question", fontsize=12, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(QS, fontsize=9); ax.set_ylim(0, 118)
    ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
    plt.tight_layout(); plt.savefig(f"imp_ragas_{metric}.png", dpi=150); plt.close()

print("Charts generated.")

# =========================================================
# WORD DOCUMENT
# =========================================================
doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)

def heading(t, level=1):
    h = doc.add_heading(t, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

def img(path):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# Title
doc.add_paragraph()
t = doc.add_heading("RAG Improvements — RAGAS Evaluation Report", level=0)
for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("ChromaDB vs FAISS vs MPNet vs Hybrid Search — Evaluated with RAGAS")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
doc.add_page_break()

# Section 1
heading("1. Overview")
doc.add_paragraph(
    "This report compares four RAG pipeline configurations evaluated using the RAGAS framework. "
    "All experiments use Groq's llama-3.3-70b for answer generation and RAGAS evaluation.")

tbl(["Experiment", "Vector Store", "Embedding", "Retrieval", "What It Tests"],
    [["Baseline", "ChromaDB", "MiniLM (384d)", "Vector similarity", "Original pipeline"],
     ["FAISS", "FAISS", "MiniLM (384d)", "Vector similarity", "Vector store speed"],
     ["MPNet", "FAISS", "MPNet (768d)", "Vector similarity", "Embedding quality"],
     ["Hybrid", "ChromaDB + BM25", "MiniLM (384d)", "BM25 + Vector merged", "Retrieval strategy"]])

# Section 2
doc.add_page_break()
heading("2. Overall Results")

tbl(["Metric", "Baseline", "FAISS", "MPNet", "Hybrid"],
    [["Faithfulness", f"{avgs['baseline'][0]*100:.1f}%", f"{avgs['faiss'][0]*100:.1f}%", f"{avgs['mpnet'][0]*100:.1f}%", f"{avgs['hybrid'][0]*100:.1f}%"],
     ["Context Recall", f"{avgs['baseline'][1]*100:.1f}%", f"{avgs['faiss'][1]*100:.1f}%", f"{avgs['mpnet'][1]*100:.1f}%", f"{avgs['hybrid'][1]*100:.1f}%"],
     ["Context Precision", f"{avgs['baseline'][2]*100:.1f}%", f"{avgs['faiss'][2]*100:.1f}%", f"{avgs['mpnet'][2]*100:.1f}%", f"{avgs['hybrid'][2]*100:.1f}%"],
     ["Retrieval Speed", f"{avgs['baseline'][3]:.0f}ms", f"{avgs['faiss'][3]:.0f}ms", f"{avgs['mpnet'][3]:.0f}ms", f"{avgs['hybrid'][3]:.0f}ms"]])

img("imp_ragas_overall.png")

heading("2.1 Retrieval Speed", level=2)
img("imp_ragas_speed.png")

# Section 3
doc.add_page_break()
heading("3. Per-Question Breakdown")
for metric, label in [("faithfulness","Faithfulness"),("context_recall","Context Recall"),("context_precision","Context Precision")]:
    heading(f"3.{['faithfulness','context_recall','context_precision'].index(metric)+1} {label}", level=2)
    tbl(["Question", "Baseline", "FAISS", "MPNet", "Hybrid"],
        [[QF[i][:45], f"{pq['baseline'][metric][i]*100:.0f}%", f"{pq['faiss'][metric][i]*100:.0f}%",
          f"{pq['mpnet'][metric][i]*100:.0f}%", f"{pq['hybrid'][metric][i]*100:.0f}%"] for i in range(5)])
    img(f"imp_ragas_{metric}.png")

# Section 4
doc.add_page_break()
heading("4. Analysis")

heading("4.1 Context Precision: The Key Differentiator", level=2)
doc.add_paragraph(
    f"RAGAS revealed that context precision improves dramatically with better methods: "
    f"Baseline {avgs['baseline'][2]*100:.0f}% → FAISS {avgs['faiss'][2]*100:.0f}% → "
    f"MPNet {avgs['mpnet'][2]*100:.0f}% → Hybrid {avgs['hybrid'][2]*100:.0f}%. "
    "This means the improved methods retrieve more relevant chunks and less noise.")

heading("4.2 Faithfulness: Hybrid Achieves Perfect Score", level=2)
doc.add_paragraph(
    f"Hybrid search achieved {avgs['hybrid'][0]*100:.0f}% faithfulness — the LLM never hallucinated "
    "beyond the retrieved context. By combining BM25 keyword matching with vector similarity, "
    "hybrid search provides the most relevant context, making it easier for the LLM to stay faithful.")

heading("4.3 MPNet Embeddings Improve Precision", level=2)
doc.add_paragraph(
    f"Upgrading from MiniLM (384 dim) to MPNet (768 dim) improved context precision from "
    f"{avgs['faiss'][2]*100:.0f}% to {avgs['mpnet'][2]*100:.0f}%. Richer embeddings capture more "
    "semantic nuance, leading to more precise document retrieval.")

heading("4.4 Speed: All Methods Comparable", level=2)
doc.add_paragraph(
    "Retrieval speeds were similar across all methods (~1,000-1,200ms). The Groq API call "
    "dominates the total time, making vector store differences negligible at this scale.")

# Section 5
heading("5. Recommended Configuration")
tbl(["Component", "Choice", "RAGAS Evidence"],
    [["Vector Store", "FAISS", f"Same recall as ChromaDB, better precision ({avgs['faiss'][2]*100:.0f}% vs {avgs['baseline'][2]*100:.0f}%)"],
     ["Embedding", "MPNet (768d)", f"Perfect context precision ({avgs['mpnet'][2]*100:.0f}%)"],
     ["Retrieval", "Hybrid (BM25+Vector)", f"Perfect faithfulness ({avgs['hybrid'][0]*100:.0f}%) and precision ({avgs['hybrid'][2]*100:.0f}%)"],
     ["Chunking", "QA-Pair", "Best across all RAGAS metrics (from chunking comparison)"]])

heading("6. Key Takeaways")
doc.add_paragraph("RAGAS context precision is the most sensitive metric to pipeline improvements — "
    "it reveals retrieval noise that faithfulness and recall don't capture.", style="List Number")
doc.add_paragraph("Hybrid search (BM25 + vector) achieves the best overall quality with perfect "
    "faithfulness and context precision.", style="List Number")
doc.add_paragraph("MPNet embeddings provide measurably better retrieval precision than MiniLM, "
    "confirming that embedding quality matters.", style="List Number")
doc.add_paragraph("All improvements are complementary — combining FAISS + MPNet + Hybrid would "
    "give the best possible configuration.", style="List Number")

doc.save("Improvements_RAGAS_Report.docx")
print("Saved: Improvements_RAGAS_Report.docx")
