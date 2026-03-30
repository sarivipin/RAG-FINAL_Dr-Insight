"""Generate 4 Word documents in ~/Downloads/ — only current project files."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.expanduser("~/Downloads")

def new_doc():
    d = Document()
    s = d.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
    return d

def heading(d, t, level=1):
    h = d.add_heading(t, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

def tbl(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    d.add_paragraph()

def p(d, text, style=None):
    d.add_paragraph(text, style=style)


# =============================================================
# DOC 1: Pipeline Execution Order
# =============================================================
def doc1():
    d = new_doc()
    t = d.add_heading("RAG Pipeline — File Execution Order", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    heading(d, "1. Data Collection & Preparation (Run Once)")

    heading(d, "Step 1: scrape_hse.py", level=2)
    p(d, "Run first. Scrapes medical condition pages from the HSE website. "
         "Collects condition names, sections, and content.\n"
         "Output: hse_conditions_long.csv\n"
         "Command: python scrape_hse.py")

    heading(d, "Step 2: convert_hse_to_qa.py", level=2)
    p(d, "Run second. Converts scraped CSV into question-answer pairs.\n"
         "Output: hse_conditions_qa.csv\n"
         "Command: python convert_hse_to_qa.py")

    heading(d, "Step 3: generate_markdown.py", level=2)
    p(d, "Run third. Converts QA CSV into structured markdown files.\n"
         "Output: 423 .md files in docs/\n"
         "Command: python generate_markdown.py")

    d.add_page_break()
    heading(d, "2. Ingestion — Building Vector Stores")

    heading(d, "Step 4: ingest.py (QA-Pair Chunking)", level=2)
    p(d, "Primary ingestion. Parses markdown into QA-pair chunks, generates embeddings, stores in ChromaDB.\n"
         "Output: db/chroma_db/\n"
         "Command: python ingest.py")

    heading(d, "Step 4a: ingest_recursive.py (Recursive Character Splitting)", level=2)
    p(d, "Alternative chunking. 800 chars with 200 overlap.\n"
         "Output: db/chroma_db_recursive/\n"
         "Command: python ingest_recursive.py")

    heading(d, "Step 4b: ingest_semantic.py (Semantic Chunking)", level=2)
    p(d, "Alternative chunking. Splits where meaning shifts via embeddings.\n"
         "Output: db/chroma_db_semantic/\n"
         "Command: python ingest_semantic.py")

    d.add_page_break()
    heading(d, "3. Running the Application")

    heading(d, "Step 5: app.py", level=2)
    p(d, "Streamlit web interface for querying medical conditions. Requires Ollama running.\n"
         "Command: streamlit run app.py")

    tbl(d, ["Supporting File", "Role"],
        [["retrieval.py", "Core RAG pipeline: retrieval, reranking, generation"],
         ["generation.py", "Interactive CLI interface (alternative to app.py)"]])

    d.add_page_break()
    heading(d, "4. Evaluation")

    heading(d, "4.1 Custom Evaluations", level=2)
    tbl(d, ["Order", "File", "What It Does"],
        [["1", "evaluation.py", "Baseline: Ollama as generator + judge"],
         ["2", "evaluation_groq.py", "Groq API as generator + judge"],
         ["3", "evaluation_fair_comparison.py", "Re-judges both answer sets with same Groq 70B judge"],
         ["4", "batch_evaluation.py", "Token-overlap metrics (no LLM judge)"]])

    heading(d, "4.2 RAGAS Evaluations", level=2)
    tbl(d, ["Order", "File", "What It Does"],
        [["5", "evaluation_ragas.py", "RAGAS standalone evaluation with Groq"],
         ["6", "evaluation_ragas_compare.py", "RAGAS vs Custom comparison on same Ollama answers"],
         ["7", "evaluation_chunking_ragas.py", "3 chunking methods evaluated with RAGAS"],
         ["8", "evaluation_improvements_ragas.py\n+ run_single_experiment_ragas.py", "ChromaDB vs FAISS vs MPNet vs Hybrid\nevaluated with RAGAS"]])

    d.add_page_break()
    heading(d, "5. Complete Execution Flow")
    p(d, "Data Pipeline:\nscrape_hse.py → convert_hse_to_qa.py → generate_markdown.py → ingest.py → app.py")
    p(d, "\nCustom Evaluation:\nevaluation.py → evaluation_groq.py → evaluation_fair_comparison.py\nbatch_evaluation.py (independent)")
    p(d, "\nRAGAS Evaluation:\nevaluation_ragas.py → evaluation_ragas_compare.py\ningest_recursive.py + ingest_semantic.py → evaluation_chunking_ragas.py\nevaluation_improvements_ragas.py (runs run_single_experiment_ragas.py per experiment)")

    d.save(os.path.join(OUT, "1_Pipeline_Execution_Order.docx"))
    print("Saved: 1_Pipeline_Execution_Order.docx")


# =============================================================
# DOC 2: Evaluations & Comparisons — Audience Guide
# =============================================================
def doc2():
    d = new_doc()
    t = d.add_heading("RAG Evaluation & Comparison — Audience Guide", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    heading(d, "1. What We Are Evaluating")
    p(d, "We built a RAG pipeline that answers medical questions by retrieving information from "
         "423 HSE medical condition documents and generating answers using an LLM.")

    heading(d, "1.1 Metrics Used", level=2)
    tbl(d, ["Metric", "What It Measures", "Used In"],
        [["Faithfulness", "Is the answer grounded in retrieved context?", "Custom + RAGAS"],
         ["Context Recall", "Did we retrieve the right documents?", "Custom + RAGAS"],
         ["Context Precision", "Were retrieved chunks actually useful?", "RAGAS only"],
         ["BLEU Score", "Word overlap with expected answer", "Custom only"]])

    d.add_page_break()
    heading(d, "2. Evaluation 1: Ollama vs Groq (Custom)")
    p(d, "Compared a small local model (Ollama 3B) with a large cloud model (Groq 70B). "
         "Each model judged its own answers.")
    tbl(d, ["Metric", "Ollama (3B)", "Groq (70B)"],
        [["Faithfulness", "56.6%", "100.0%"],
         ["Context Recall", "38.6%", "100.0%"],
         ["BLEU Score", "51.9%", "63.7%"]])
    p(d, "How to explain: \"The initial gap looked massive, but we discovered the small model "
         "was a poor judge — like having a student grade their own exam versus a professor.\"")

    d.add_page_break()
    heading(d, "3. Evaluation 2: Fair Comparison (Same Judge)")
    p(d, "Re-evaluated both answer sets using the same Groq 70B judge to isolate answer quality from judge quality.")
    tbl(d, ["Metric", "Ollama Answers (Groq-judged)", "Groq Answers (Groq-judged)"],
        [["Faithfulness", "88.0%", "100.0%"],
         ["Context Recall", "100.0%", "100.0%"],
         ["BLEU Score", "51.9%", "63.7%"]])
    p(d, "How to explain: \"When we used the same judge, Ollama's faithfulness jumped from 56.6% to 88.0%. "
         "The answers were always decent — the judge was the problem. The real gap is only 12%.\"")

    d.add_page_break()
    heading(d, "4. Evaluation 3: RAGAS vs Custom Comparison")
    p(d, "Evaluated the same Ollama answers using both our custom prompt-based approach and the "
         "industry-standard RAGAS framework.")
    tbl(d, ["Metric", "RAGAS (Groq 70B judge)", "Custom (Ollama 3B judge)"],
        [["Faithfulness", "92.5%", "56.6%"],
         ["Context Recall", "100.0%", "38.6%"],
         ["Context Precision", "93.3%", "N/A"],
         ["BLEU", "N/A", "51.9%"]])
    p(d, "How to explain: \"RAGAS uses a structured approach — it extracts individual claims from the answer "
         "and verifies each one against the context. This is more rigorous than asking an LLM to give a single "
         "score. RAGAS also adds Context Precision, which tells us if we retrieved irrelevant documents.\"")

    d.add_page_break()
    heading(d, "5. Evaluation 4: Chunking Methods (RAGAS)")
    p(d, "Tested three chunking strategies using RAGAS evaluation.")
    tbl(d, ["Method", "Faithfulness", "Context Recall", "Context Precision"],
        [["QA-Pair (Original)", "98.0%", "100.0%", "90.0%"],
         ["Recursive Character", "98.0%", "80.0%", "78.3%"],
         ["Semantic Chunking", "97.1%", "66.0%", "71.7%"]])
    p(d, "How to explain: \"Since our documents are structured as Q&A pairs, chunking by those natural "
         "boundaries works best. Generic splitting methods break questions apart from their answers, "
         "hurting retrieval. Chunking strategy should match document structure.\"")

    d.add_page_break()
    heading(d, "6. Evaluation 5: RAG Improvements (RAGAS)")
    p(d, "Tested four pipeline configurations using RAGAS.")
    tbl(d, ["Method", "Faithfulness", "Context Recall", "Context Precision", "Speed"],
        [["Baseline (ChromaDB+MiniLM)", "98.0%", "100.0%", "88.3%", "1,081ms"],
         ["FAISS+MiniLM", "95.0%", "100.0%", "98.3%", "1,093ms"],
         ["FAISS+MPNet", "95.0%", "100.0%", "100.0%", "1,088ms"],
         ["Hybrid (BM25+Vector)", "100.0%", "98.2%", "100.0%", "1,190ms"]])
    p(d, "How to explain: \"Each improvement targets a different bottleneck. FAISS improves retrieval precision. "
         "MPNet embeddings capture more meaning. Hybrid search combines keyword and semantic matching for "
         "perfect faithfulness and precision. RAGAS's context precision metric revealed improvements "
         "that simpler evaluation methods couldn't detect.\"")

    d.add_page_break()
    heading(d, "7. What I Chose and Why")
    tbl(d, ["Component", "Choice", "Why"],
        [["Chunking", "QA-Pair", "Documents are structured as Q&A — matches naturally"],
         ["Vector Store", "FAISS", "Better context precision than ChromaDB (98.3% vs 88.3%)"],
         ["Embedding", "all-mpnet-base-v2", "Perfect context precision (100%) with richer 768-dim embeddings"],
         ["Retrieval", "Hybrid (BM25+Vector)", "Perfect faithfulness (100%) and precision (100%)"],
         ["LLM", "Groq llama-3.3-70b", "Best answer quality, fast API"],
         ["Evaluation", "RAGAS framework", "Industry-standard, structured decomposition, more rigorous"]])

    d.save(os.path.join(OUT, "2_Evaluations_Audience_Guide.docx"))
    print("Saved: 2_Evaluations_Audience_Guide.docx")


# =============================================================
# DOC 3: Python File Explanations
# =============================================================
def doc3():
    d = new_doc()
    t = d.add_heading("Python File Explanations", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    files = [
        ("scrape_hse.py", "Data Collection — Web Scraper",
         "Scrapes medical condition pages from https://www2.hse.ie/conditions/. Visits each condition page "
         "to extract structured content (condition name, sections, text). Includes 1-second delay between requests.",
         "Key functions: extract_condition_links(), scrape_condition(), normalize_heading(), extract_section_content()\n"
         "Output: hse_conditions_long.csv"),

        ("convert_hse_to_qa.py", "Data Transformation — CSV to QA Pairs",
         "Converts raw scraped data into question-answer format using predefined templates. "
         "E.g., 'symptoms' section → 'What are the symptoms of [condition]?'",
         "Key logic: SECTION_QUESTION_MAP templates, deduplication, junk filtering\n"
         "Output: hse_conditions_qa.csv"),

        ("generate_markdown.py", "Data Formatting — QA CSV to Markdown",
         "Converts QA CSV into structured markdown files, one per condition. "
         "Uses ## Question / ### Answer / **Section:** / **Source:** format.",
         "Output: 423 .md files in docs/"),

        ("ingest.py", "Ingestion — QA-Pair Chunking (Primary)",
         "Parses markdown files into QA-pair chunks, generates embeddings (all-MiniLM-L6-v2), stores in ChromaDB. "
         "Splits long answers at paragraph/sentence boundaries (800 char max).",
         "Key functions: split_into_question_blocks(), parse_question_block(), split_long_answer(), make_question_hash()\n"
         "Output: db/chroma_db/ (~1,950 chunks)"),

        ("ingest_recursive.py", "Ingestion — Recursive Character Splitting",
         "Alternative ingestion using RecursiveCharacterTextSplitter. Splits at natural boundaries "
         "(paragraphs → sentences → words) with 800 char chunks and 200 char overlap.",
         "Output: db/chroma_db_recursive/ (2,801 chunks)"),

        ("ingest_semantic.py", "Ingestion — Semantic Chunking",
         "Alternative ingestion using SemanticChunker. Detects meaning shifts via embeddings "
         "and splits at semantic boundaries. Uses percentile threshold of 70.",
         "Output: db/chroma_db_semantic/ (4,248 chunks)"),

        ("retrieval.py", "Core RAG Pipeline",
         "The heart of the system. Handles query processing, disease detection, document retrieval "
         "(similarity + MMR), metadata boosting, cross-encoder reranking, question group merging, and answer generation.",
         "Key functions: ask() (main entry), extract_disease_from_query(), detect_section(), "
         "similarity_retrieve(), mmr_retrieve(), rerank_documents(), select_top_question_groups()\n"
         "Uses Ollama llama3.2 for generation by default."),

        ("generation.py", "Interactive CLI Interface",
         "Command-line interface for asking questions interactively. Saves Q&A history to JSON.",
         "Uses retrieval.py's ask() function internally."),

        ("app.py", "Streamlit Web Application",
         "Main user-facing web interface. Disease dropdown, question suggestions, custom queries, "
         "configurable retrieval method (similarity/MMR/hybrid), confidence scoring, debug mode.",
         "Command: streamlit run app.py"),
    ]

    for fname, title, desc, details in files:
        heading(d, f"{fname} — {title}")
        p(d, desc)
        p(d, details)
        d.add_paragraph()

    d.add_page_break()
    heading(d, "Evaluation Scripts", level=0)

    eval_files = [
        ("evaluation.py", "Baseline Evaluation (Ollama)",
         "Uses Ollama llama3.2 (3B) as both generator and judge. Implements custom faithfulness, "
         "context recall (LLM-judged), and BLEU score (algorithmic).",
         "Output: ragas_results.csv, ragas_details.json"),

        ("evaluation_groq.py", "API-Based Evaluation (Groq)",
         "Same logic as evaluation.py but uses Groq llama-3.3-70b (70B) via API. "
         "Includes retry logic with exponential backoff for rate limiting. "
         "Patches retrieval.load_llm to swap Ollama for Groq.",
         "Output: ragas_results_groq.csv, ragas_details_groq.json"),

        ("evaluation_fair_comparison.py", "Fair Comparison (Same Judge)",
         "Loads saved answers from both Ollama and Groq runs. Re-evaluates ALL answers using "
         "the same Groq 70B judge. Isolates answer quality from judge quality.",
         "Output: ragas_results_fair_comparison.csv, Fair_Comparison_Report.docx"),

        ("batch_evaluation.py", "Token-Overlap Evaluation",
         "Fast algorithmic evaluation — no LLM judge needed. Measures faithfulness (answer-context overlap), "
         "coverage (context-answer overlap), and relevance (query-answer overlap).",
         "Output: batch_results.csv, batch_details.json"),

        ("evaluation_ragas.py", "RAGAS Standalone Evaluation",
         "Uses the RAGAS framework with Groq 70B as judge. Evaluates faithfulness, context recall, "
         "and context precision using RAGAS's structured claim decomposition approach.",
         "Output: ragas_lib_results.csv, ragas_lib_details.json"),

        ("evaluation_ragas_compare.py", "RAGAS vs Custom Comparison",
         "Loads saved Ollama answers and evaluates them with RAGAS. Compares RAGAS scores against "
         "the original custom evaluation scores to show the difference in methodology.",
         "Output: ragas_vs_custom_results.csv, RAGAS_vs_Custom_Comparison_Report.docx"),

        ("evaluation_chunking_ragas.py", "Chunking Methods — RAGAS Evaluation",
         "Evaluates all 3 chunking methods (QA-pair, recursive, semantic) using RAGAS. "
         "Each method uses its own vector store. Standalone retrieval pipeline.",
         "Output: chunking_ragas_results.json, Chunking_RAGAS_Comparison_Report.docx"),

        ("evaluation_improvements_ragas.py + run_single_experiment_ragas.py", "RAG Improvements — RAGAS",
         "Tests 4 configurations: Baseline (ChromaDB), FAISS, FAISS+MPNet, Hybrid (BM25+Vector). "
         "Each experiment runs in a separate subprocess. All evaluated with RAGAS.",
         "Output: improvements_ragas_results.json, Improvements_RAGAS_Report.docx"),
    ]

    for fname, title, desc, details in eval_files:
        heading(d, f"{fname} — {title}")
        p(d, desc)
        p(d, details)
        d.add_paragraph()

    d.save(os.path.join(OUT, "3_Python_File_Explanations.docx"))
    print("Saved: 3_Python_File_Explanations.docx")


# =============================================================
# DOC 4: Evaluation Challenges & How I Overcame Them
# =============================================================
def doc4():
    d = new_doc()
    t = d.add_heading("Evaluation Challenges & How I Overcame Them", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A journey through building and evaluating a Medical RAG Pipeline")
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55); r.italic = True
    d.add_page_break()

    # Challenge 1
    heading(d, "Challenge 1: Getting the Evaluation Pipeline Running")
    p(d, "When I first tried to run evaluation.py, it wouldn't start. The script depended on Ollama "
         "running as a local server, but the service wasn't active. I got connection errors every time.")
    heading(d, "How I Solved It", level=2)
    p(d, "I realised Ollama needs to run as a background service. I started it with 'ollama serve' "
         "and verified the llama3.2 model was available. Once running, evaluation.py executed successfully: "
         "Faithfulness 56.6%, Context Recall 38.6%, BLEU 51.9%.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Local LLM setups require the inference server to be running — a common gotcha when "
         "switching between local and API-based models.")

    d.add_page_break()
    # Challenge 2
    heading(d, "Challenge 2: Gemini API Rate Limits")
    p(d, "I wanted to compare with a cloud LLM. I created an evaluation script using Google's Gemini API "
         "with a free-tier key. It processed the first question, then crashed with 429 RESOURCE_EXHAUSTED.")
    heading(d, "What I Tried", level=2)
    p(d, "I obtained a new Gemini key, but it hit the same limits. I added retry logic with exponential "
         "backoff (30s, 60s, 120s, 240s), but the daily quota was exhausted entirely.")
    heading(d, "How I Solved It", level=2)
    p(d, "I researched alternatives and found Groq, which offers a generous free tier with ~14,400 "
         "requests/day. I installed langchain-groq, created evaluation_groq.py, and it ran perfectly "
         "on the first attempt — all 15 LLM calls completed without rate limiting.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Free-tier API keys have strict quotas. Always check rate limits before committing to a provider. "
         "Groq's free tier was ideal for evaluation workloads.")

    d.add_page_break()
    # Challenge 3
    heading(d, "Challenge 3: Misleading Evaluation Results")
    p(d, "Comparing Ollama and Groq results showed an enormous gap: 56.6% vs 100% faithfulness. "
         "But BLEU scores only improved modestly (51.9% → 63.7%). Something didn't add up.")
    heading(d, "How I Investigated", level=2)
    p(d, "I realised each model was judging its own answers. The 3B model was both generating AND "
         "evaluating — a small model is a weak evaluator that gives inconsistent scores.")
    heading(d, "How I Solved It", level=2)
    p(d, "I created evaluation_fair_comparison.py that loads saved answers from both models and "
         "re-evaluates ALL of them using the same Groq 70B judge. Ollama's faithfulness jumped from "
         "56.6% to 88.0% — the answers were always decent, the judge was the problem.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Never use the same model as both generator and judge. Always use a strong, separate model "
         "for evaluation. BLEU scores serve as a useful sanity check with no LLM bias.")

    d.add_page_break()
    # Challenge 4
    heading(d, "Challenge 4: Adopting RAGAS for Industry-Standard Evaluation")
    p(d, "I wanted to validate my custom evaluation results using the RAGAS framework — the "
         "industry standard for RAG evaluation. But RAGAS makes many more LLM calls than custom "
         "evaluation (~50 calls vs ~15), which exhausted my free Groq API quota.")
    heading(d, "How I Solved It", level=2)
    p(d, "I upgraded to Groq's Developer tier (pay-as-you-go, a few cents per evaluation run). "
         "RAGAS then ran successfully and provided more granular insights. I compared RAGAS scores "
         "against my custom scores on the same Ollama answers: RAGAS gave 92.5% faithfulness vs "
         "custom's 56.6%, confirming the weak judge was the issue, not the answers.")
    p(d, "RAGAS also provided Context Precision (93.3%) — a metric my custom evaluation didn't have — "
         "which revealed retrieval noise that simpler methods couldn't detect.")
    heading(d, "Lesson Learned", level=2)
    p(d, "RAGAS's structured decomposition (extracting claims, verifying each one) is more rigorous "
         "than single-prompt scoring. For production systems, RAGAS is the recommended approach.")

    d.add_page_break()
    # Challenge 5
    heading(d, "Challenge 5: Choosing the Right Chunking Strategy")
    p(d, "I needed to test whether my QA-pair chunking was actually the best approach, or if "
         "standard methods like recursive splitting or semantic chunking would perform better.")
    heading(d, "How I Solved It", level=2)
    p(d, "I implemented three chunking methods, each with its own ingestion script and vector store. "
         "I evaluated all three using RAGAS to get rigorous, comparable metrics.")
    p(d, "Results clearly showed QA-pair chunking winning: 98% faithfulness, 100% context recall, "
         "90% context precision. Recursive splitting dropped to 80% recall, and semantic chunking "
         "to 66% recall. The documents are structured as Q&A pairs, so chunking by those boundaries "
         "preserves semantic completeness.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Chunking strategy should match document structure. More chunks doesn't mean better retrieval — "
         "semantic chunking created 2x more chunks but scored lowest.")

    d.add_page_break()
    # Challenge 6
    heading(d, "Challenge 6: Memory Crashes with FAISS Experiments")
    p(d, "When testing FAISS, MPNet, and hybrid search in a single Python process, the script crashed "
         "with a segmentation fault. Loading ChromaDB, FAISS, two embedding models, and the reranker "
         "simultaneously exceeded available memory.")
    heading(d, "How I Solved It", level=2)
    p(d, "I restructured into two files: evaluation_improvements_ragas.py (orchestrator) launches "
         "each experiment via run_single_experiment_ragas.py as a separate subprocess. Each experiment "
         "runs in its own process with fresh memory, then writes results to JSON. The orchestrator "
         "collects all results and runs RAGAS evaluation.")
    heading(d, "Lesson Learned", level=2)
    p(d, "When working with multiple ML models and vector stores, memory management matters. "
         "Subprocess isolation is a clean solution for avoiding library conflicts.")

    d.add_page_break()
    # Challenge 7
    heading(d, "Challenge 7: Comparing Different Improvements Fairly")
    p(d, "I tested four pipeline configurations (ChromaDB, FAISS, MPNet, Hybrid) and needed to ensure "
         "each comparison was fair — changing only one variable at a time.")
    heading(d, "How I Solved It", level=2)
    p(d, "I used RAGAS to evaluate all four configurations with the same judge, same questions, "
         "and same ground truth. RAGAS revealed that context precision was the key differentiator: "
         "Baseline 88.3% → FAISS 98.3% → MPNet 100% → Hybrid 100%. "
         "Hybrid search achieved perfect faithfulness (100%) and precision (100%).")
    heading(d, "Lesson Learned", level=2)
    p(d, "RAGAS's context precision metric was crucial — it detected retrieval noise that faithfulness "
         "and recall couldn't capture. Each improvement targets a different bottleneck, and they're "
         "complementary: FAISS for speed, MPNet for embedding quality, hybrid for coverage.")

    d.add_page_break()
    heading(d, "Summary: The Journey")
    p(d, "Building the RAG pipeline was straightforward. Evaluating it properly was the real challenge. "
         "Along the way, I learned:")
    p(d, "1. Local LLM setup requires attention to service management", style="List Number")
    p(d, "2. Free API tiers have real limits — always have a backup plan", style="List Number")
    p(d, "3. Self-evaluation is unreliable — use a separate, strong judge model", style="List Number")
    p(d, "4. RAGAS provides more rigorous evaluation than custom prompt-based scoring", style="List Number")
    p(d, "5. Chunking strategy should match document structure", style="List Number")
    p(d, "6. Memory management matters when combining multiple ML models", style="List Number")
    p(d, "7. RAGAS context precision reveals insights that simpler metrics miss", style="List Number")
    p(d, "")
    p(d, "The final recommended pipeline — QA-pair chunking, FAISS, MPNet embeddings, hybrid search, "
         "and Groq's 70B model — represents the best configuration found through systematic RAGAS-evaluated experimentation.")

    d.save(os.path.join(OUT, "4_Evaluation_Challenges_Story.docx"))
    print("Saved: 4_Evaluation_Challenges_Story.docx")


# =============================================================
# RUN ALL
# =============================================================
if __name__ == "__main__":
    doc1()
    doc2()
    doc3()
    doc4()
    print(f"\nAll 4 documents saved to: {OUT}")
