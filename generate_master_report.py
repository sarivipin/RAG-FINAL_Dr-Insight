"""
Master Report: All evaluations, comparisons, and chunking analysis in one document.
"""
import json
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# =========================================================
# DATA
# =========================================================
QUESTIONS_FULL = [
    "Symptoms of bacterial vaginosis",
    "Causes of varicose veins",
    "Preventing melanoma",
    "Living with type 1 diabetes",
    "MRSA",
]
QUESTIONS_SHORT = ["Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention", "Type 1\nDiabetes", "MRSA"]

# Original self-judged
orig_ollama = {"faith": [0.80,0.80,0.33,0.50,0.40], "recall": [0.40,0.40,0.40,0.33,0.40], "bleu": [0.911,0.459,0.785,0.096,0.342]}
orig_groq   = {"faith": [1.0,1.0,1.0,1.0,1.0],      "recall": [1.0,1.0,1.0,1.0,1.0],      "bleu": [0.868,0.361,0.969,0.200,0.786]}

# Fair comparison (both judged by Groq 70B)
fair_ollama = {"faith": [1.0,1.0,1.0,0.8,0.6], "recall": [1.0,1.0,1.0,1.0,1.0], "bleu": [0.911,0.459,0.785,0.096,0.342]}
fair_groq   = {"faith": [1.0,1.0,1.0,1.0,1.0], "recall": [1.0,1.0,1.0,1.0,1.0], "bleu": [0.868,0.361,0.969,0.200,0.786]}

# Batch (token-overlap)
batch = {"faith": [0.95,0.935,0.96,0.319,0.667], "coverage": [0.266,0.218,0.289,0.344,0.352], "relevance": [0.571,0.250,0.714,0.800,0.545]}

# Chunking comparison
with open("chunking_comparison_results.json") as f:
    chunk_data = json.load(f)
chunk_methods = ["qa_pair", "recursive", "semantic"]
chunk_labels = [chunk_data[m]["label"] for m in chunk_methods]
chunk_avgs = {m: [chunk_data[m]["avg_faithfulness"], chunk_data[m]["avg_context_recall"], chunk_data[m]["avg_bleu"]] for m in chunk_methods}
chunk_pq = {}
for m in chunk_methods:
    chunk_pq[m] = {
        "faith": [r["faithfulness"] for r in chunk_data[m]["per_question"]],
        "recall": [r["context_recall"] for r in chunk_data[m]["per_question"]],
        "bleu": [r["bleu"] for r in chunk_data[m]["per_question"]],
    }

avg = lambda lst: sum(lst)/len(lst)
C = {"ollama": "#4A90D9", "groq": "#E8833A", "batch": "#2ECC71", "rec": "#2ECC71", "sem": "#9B59B6", "qa": "#4A90D9"}

# =========================================================
# CHART HELPERS
# =========================================================
def save_grouped_bar(filename, title, x_labels, groups, group_labels, colors, figsize=(9,5.5)):
    fig, ax = plt.subplots(figsize=figsize)
    x = np.arange(len(x_labels))
    n = len(groups)
    w = 0.7 / n
    for i, (vals, lbl, col) in enumerate(zip(groups, group_labels, colors)):
        offset = (i - (n-1)/2) * w
        bars = ax.bar(x + offset, [v*100 for v in vals], w, label=lbl, color=col)
        for b in bars:
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+1, f"{b.get_height():.0f}%", ha="center", fontsize=7)
    ax.set_ylabel("Score (%)"); ax.set_title(title, fontsize=12, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(x_labels, fontsize=9)
    ax.set_ylim(0, 118); ax.legend(fontsize=8); ax.grid(axis="y", alpha=0.3)
    plt.tight_layout(); plt.savefig(filename, dpi=150); plt.close()

# Chart 1: Original comparison overall
save_grouped_bar("m_orig_overall.png", "Original Evaluation: Overall Averages",
    ["Faithfulness", "Context Recall", "BLEU"],
    [[avg(orig_ollama["faith"]), avg(orig_ollama["recall"]), avg(orig_ollama["bleu"])],
     [avg(orig_groq["faith"]), avg(orig_groq["recall"]), avg(orig_groq["bleu"])]],
    ["Ollama (self-judged)", "Groq (self-judged)"], [C["ollama"], C["groq"]])

# Chart 2: Fair comparison overall
save_grouped_bar("m_fair_overall.png", "Fair Comparison: Overall (Same Groq 70B Judge)",
    ["Faithfulness", "Context Recall", "BLEU"],
    [[avg(fair_ollama["faith"]), avg(fair_ollama["recall"]), avg(fair_ollama["bleu"])],
     [avg(fair_groq["faith"]), avg(fair_groq["recall"]), avg(fair_groq["bleu"])]],
    ["Ollama answers", "Groq answers"], [C["ollama"], C["groq"]])

# Chart 3: Ollama self-judged vs Groq-judged
save_grouped_bar("m_judge_impact.png", "Impact of Judge Quality on Ollama Scores",
    ["Faithfulness", "Context Recall", "BLEU"],
    [[avg(orig_ollama["faith"]), avg(orig_ollama["recall"]), avg(orig_ollama["bleu"])],
     [avg(fair_ollama["faith"]), avg(fair_ollama["recall"]), avg(fair_ollama["bleu"])]],
    ["Ollama (self-judged by 3B)", "Ollama (judged by Groq 70B)"], ["#85C1E9", C["ollama"]])

# Chart 4: Batch vs Fair faithfulness per question
save_grouped_bar("m_batch_vs_fair.png", "Faithfulness: Batch (Token-Overlap) vs Fair (LLM-Judged)",
    QUESTIONS_SHORT,
    [batch["faith"], fair_ollama["faith"], fair_groq["faith"]],
    ["Batch (token-overlap)", "Fair: Ollama", "Fair: Groq"],
    [C["batch"], C["ollama"], C["groq"]], figsize=(11,5.5))

# Chart 5: Chunking overall
save_grouped_bar("m_chunk_overall.png", "Chunking Methods: Overall Averages",
    ["Faithfulness", "Context Recall", "BLEU"],
    [[chunk_avgs[m][i] for i in range(3)] for m in chunk_methods],
    chunk_labels, [C["qa"], C["rec"], C["sem"]])

# Chart 6: Chunking per-question faithfulness
save_grouped_bar("m_chunk_faith.png", "Chunking: Faithfulness per Question",
    QUESTIONS_SHORT,
    [chunk_pq[m]["faith"] for m in chunk_methods],
    chunk_labels, [C["qa"], C["rec"], C["sem"]], figsize=(11,5.5))

# Chart 7: Chunking per-question recall
save_grouped_bar("m_chunk_recall.png", "Chunking: Context Recall per Question",
    QUESTIONS_SHORT,
    [chunk_pq[m]["recall"] for m in chunk_methods],
    chunk_labels, [C["qa"], C["rec"], C["sem"]], figsize=(11,5.5))

# Chart 8: Chunking per-question BLEU
save_grouped_bar("m_chunk_bleu.png", "Chunking: BLEU Score per Question",
    QUESTIONS_SHORT,
    [chunk_pq[m]["bleu"] for m in chunk_methods],
    chunk_labels, [C["qa"], C["rec"], C["sem"]], figsize=(11,5.5))

print("All charts generated.")

# =========================================================
# WORD DOCUMENT
# =========================================================
doc = Document()
style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

def chart(path):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def p(text, style_name=None):
    doc.add_paragraph(text, style=style_name)

# ---- TITLE ----
doc.add_paragraph()
t = doc.add_heading("RAG Pipeline — Master Evaluation Report", level=0)
for run in t.runs: run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Complete evaluation across models, judges, metrics, and chunking strategies")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55); r.italic = True
doc.add_page_break()

# =========================================================
# SECTION 1: PROJECT & EVALUATION OVERVIEW
# =========================================================
heading("1. Project Overview", level=1)

p("This report consolidates all evaluations performed on the Medical RAG pipeline. "
  "The pipeline retrieves information from 423 HSE medical condition documents and generates "
  "answers using an LLM. Five test questions from questions.txt are used throughout.")

heading("1.1 Data Sources", level=2)
table(
    ["Component", "Source", "Details"],
    [
        ["Questions", "questions.txt", "5 medical questions"],
        ["Ground Truth", "docs/ folder (423 .md files)", "1,942 QA pairs extracted from markdown"],
        ["Vector Store", "db/chroma_db (+ variants)", "ChromaDB with sentence-transformer embeddings"],
    ],
)

heading("1.2 Evaluation Files", level=2)
p("Five separate evaluation scripts were created, each serving a different purpose:")

table(
    ["File", "Purpose", "Answer Generator", "Evaluation Judge", "Metrics"],
    [
        ["evaluation.py", "Original baseline evaluation", "Ollama\n(llama3.2, 3B)", "Ollama\n(llama3.2, 3B)", "LLM-judged Faithfulness\nLLM-judged Context Recall\nBLEU Score"],
        ["evaluation_groq.py", "API-based evaluation", "Groq API\n(llama-3.3-70b, 70B)", "Groq API\n(llama-3.3-70b, 70B)", "LLM-judged Faithfulness\nLLM-judged Context Recall\nBLEU Score"],
        ["evaluation_fair_comparison.py", "Fair comparison with\nsame judge for both", "N/A\n(uses saved answers)", "Groq API\n(llama-3.3-70b, 70B)", "LLM-judged Faithfulness\nLLM-judged Context Recall\nBLEU Score"],
        ["batch_evaluation.py", "Fast algorithmic evaluation", "Ollama\n(llama3.2, 3B)", "None\n(algorithmic only)", "Token-overlap Faithfulness\nToken-overlap Coverage\nToken-overlap Relevance"],
        ["evaluation_chunking_comparison.py", "Compare 3 chunking methods", "Groq API\n(llama-3.3-70b, 70B)", "Groq API\n(llama-3.3-70b, 70B)", "LLM-judged Faithfulness\nLLM-judged Context Recall\nBLEU Score"],
    ],
)

heading("1.3 Metrics Explained", level=2)
table(
    ["Metric", "Type", "What It Measures"],
    [
        ["Faithfulness\n(LLM-judged)", "LLM as judge", "Is the answer grounded in the retrieved context?\nScored 0.0–1.0 by the judge LLM"],
        ["Context Recall\n(LLM-judged)", "LLM as judge", "Does the retrieved context cover the ground-truth reference?\nScored 0.0–1.0 by the judge LLM"],
        ["BLEU Score", "Algorithmic", "Word overlap between generated answer and reference answer.\nNo LLM involved — purely deterministic"],
        ["Faithfulness\n(token-overlap)", "Algorithmic", "Fraction of answer tokens found in context.\nUsed in batch_evaluation.py only"],
        ["Coverage\n(token-overlap)", "Algorithmic", "Fraction of context tokens found in answer.\nUsed in batch_evaluation.py only"],
        ["Relevance\n(token-overlap)", "Algorithmic", "Fraction of query tokens found in answer.\nUsed in batch_evaluation.py only"],
    ],
)

# =========================================================
# SECTION 2: ORIGINAL COMPARISON (Ollama vs Groq, self-judged)
# =========================================================
doc.add_page_break()
heading("2. Original Evaluation: Ollama vs Groq (Self-Judged)", level=1)

p("In this evaluation, each model served as both the answer generator and the evaluation judge. "
  "Ollama's llama3.2 (3B) judged its own answers, and Groq's llama-3.3-70b (70B) judged its own.")

table(
    ["Metric", "Ollama (llama3.2)", "Groq (llama-3.3-70b)", "Difference"],
    [
        ["Faithfulness", "56.6%", "100.0%", "+43.4%"],
        ["Context Recall", "38.6%", "100.0%", "+61.4%"],
        ["BLEU Score", "51.9%", "63.7%", "+11.8%"],
    ],
)
chart("m_orig_overall.png")

heading("2.1 Per-Question Scores", level=2)
table(
    ["Question", "Ollama Faith.", "Groq Faith.", "Ollama Recall", "Groq Recall", "Ollama BLEU", "Groq BLEU"],
    [[QUESTIONS_FULL[i],
      f"{orig_ollama['faith'][i]*100:.0f}%", f"{orig_groq['faith'][i]*100:.0f}%",
      f"{orig_ollama['recall'][i]*100:.0f}%", f"{orig_groq['recall'][i]*100:.0f}%",
      f"{orig_ollama['bleu'][i]*100:.1f}%", f"{orig_groq['bleu'][i]*100:.1f}%"] for i in range(5)],
)

# =========================================================
# SECTION 3: FAIR COMPARISON (same Groq judge)
# =========================================================
doc.add_page_break()
heading("3. Fair Comparison: Same Judge (Groq 70B) for Both", level=1)

p("To isolate answer quality from judge quality, both sets of answers were re-evaluated "
  "using the same Groq 70B judge. The original generated answers were kept unchanged.")

table(
    ["Metric", "Ollama Answers", "Groq Answers", "Difference"],
    [
        ["Faithfulness", f"{avg(fair_ollama['faith'])*100:.1f}%", f"{avg(fair_groq['faith'])*100:.1f}%",
         f"{(avg(fair_groq['faith'])-avg(fair_ollama['faith']))*100:+.1f}%"],
        ["Context Recall", f"{avg(fair_ollama['recall'])*100:.1f}%", f"{avg(fair_groq['recall'])*100:.1f}%",
         f"{(avg(fair_groq['recall'])-avg(fair_ollama['recall']))*100:+.1f}%"],
        ["BLEU Score", f"{avg(fair_ollama['bleu'])*100:.1f}%", f"{avg(fair_groq['bleu'])*100:.1f}%",
         f"{(avg(fair_groq['bleu'])-avg(fair_ollama['bleu']))*100:+.1f}%"],
    ],
)
chart("m_fair_overall.png")

heading("3.1 Impact of Judge Quality", level=2)
p("This table shows how Ollama's scores changed when re-evaluated by the Groq 70B judge:")
table(
    ["Metric", "Ollama (self-judged by 3B)", "Ollama (judged by Groq 70B)", "Change"],
    [
        ["Faithfulness", "56.6%", f"{avg(fair_ollama['faith'])*100:.1f}%", f"{avg(fair_ollama['faith'])*100-56.6:+.1f}%"],
        ["Context Recall", "38.6%", f"{avg(fair_ollama['recall'])*100:.1f}%", f"{avg(fair_ollama['recall'])*100-38.6:+.1f}%"],
        ["BLEU Score", "51.9%", f"{avg(fair_ollama['bleu'])*100:.1f}%", "0.0% (algorithmic)"],
    ],
)
chart("m_judge_impact.png")

# =========================================================
# SECTION 4: BATCH EVALUATION (token-overlap)
# =========================================================
doc.add_page_break()
heading("4. Batch Evaluation: Token-Overlap Metrics", level=1)

p("The batch evaluation uses simple token-overlap metrics — no LLM is involved in scoring. "
  "This provides fast, deterministic, and free evaluation but only captures surface-level matching.")

table(
    ["Metric", "Average Score", "How It Works"],
    [
        ["Faithfulness", f"{avg(batch['faith'])*100:.1f}%", "answer tokens ∩ context tokens / answer tokens"],
        ["Coverage", f"{avg(batch['coverage'])*100:.1f}%", "answer tokens ∩ context tokens / context tokens"],
        ["Relevance", f"{avg(batch['relevance'])*100:.1f}%", "query tokens ∩ answer tokens / query tokens"],
    ],
)

heading("4.1 Batch vs LLM-Judged Faithfulness", level=2)
p("Comparing the same Ollama-generated answers scored by token-overlap vs LLM judge:")
table(
    ["Question", "Batch\n(token-overlap)", "Fair\n(Groq 70B judge)"],
    [[QUESTIONS_FULL[i], f"{batch['faith'][i]*100:.1f}%", f"{fair_ollama['faith'][i]*100:.1f}%"] for i in range(5)],
)
chart("m_batch_vs_fair.png")

# =========================================================
# SECTION 5: CHUNKING COMPARISON
# =========================================================
doc.add_page_break()
heading("5. Chunking Method Comparison", level=1)

p("Three chunking strategies were tested. All used Groq 70B as both generator and judge. "
  "Each method writes to its own vector store — the original DB was not modified.")

heading("5.1 Chunking Methods", level=2)
table(
    ["", "QA-Pair (Original)", "Recursive Character", "Semantic Chunking"],
    [
        ["File", "ingest.py", "ingest_recursive.py", "ingest_semantic.py"],
        ["Method", "Split by ## Question blocks", "RecursiveCharacterTextSplitter\n800 chars, 200 overlap", "SemanticChunker\n(embedding similarity)"],
        ["Chunks", "~1,950", "2,801", "4,248"],
        ["Overlap", "None", "200 characters", "None"],
        ["DB Path", "db/chroma_db", "db/chroma_db_recursive", "db/chroma_db_semantic"],
    ],
)

heading("5.2 Overall Results", level=2)
table(
    ["Metric", "QA-Pair", "Recursive", "Semantic"],
    [
        ["Faithfulness", f"{chunk_avgs['qa_pair'][0]*100:.1f}%", f"{chunk_avgs['recursive'][0]*100:.1f}%", f"{chunk_avgs['semantic'][0]*100:.1f}%"],
        ["Context Recall", f"{chunk_avgs['qa_pair'][1]*100:.1f}%", f"{chunk_avgs['recursive'][1]*100:.1f}%", f"{chunk_avgs['semantic'][1]*100:.1f}%"],
        ["BLEU Score", f"{chunk_avgs['qa_pair'][2]*100:.1f}%", f"{chunk_avgs['recursive'][2]*100:.1f}%", f"{chunk_avgs['semantic'][2]*100:.1f}%"],
    ],
)
chart("m_chunk_overall.png")

heading("5.3 Per-Question: Faithfulness", level=2)
table(
    ["Question", "QA-Pair", "Recursive", "Semantic"],
    [[QUESTIONS_FULL[i], f"{chunk_pq['qa_pair']['faith'][i]*100:.0f}%",
      f"{chunk_pq['recursive']['faith'][i]*100:.0f}%", f"{chunk_pq['semantic']['faith'][i]*100:.0f}%"] for i in range(5)],
)
chart("m_chunk_faith.png")

heading("5.4 Per-Question: Context Recall", level=2)
table(
    ["Question", "QA-Pair", "Recursive", "Semantic"],
    [[QUESTIONS_FULL[i], f"{chunk_pq['qa_pair']['recall'][i]*100:.0f}%",
      f"{chunk_pq['recursive']['recall'][i]*100:.0f}%", f"{chunk_pq['semantic']['recall'][i]*100:.0f}%"] for i in range(5)],
)
chart("m_chunk_recall.png")

heading("5.5 Per-Question: BLEU Score", level=2)
table(
    ["Question", "QA-Pair", "Recursive", "Semantic"],
    [[QUESTIONS_FULL[i], f"{chunk_pq['qa_pair']['bleu'][i]*100:.1f}%",
      f"{chunk_pq['recursive']['bleu'][i]*100:.1f}%", f"{chunk_pq['semantic']['bleu'][i]*100:.1f}%"] for i in range(5)],
)
chart("m_chunk_bleu.png")

# =========================================================
# SECTION 6: ANALYSIS
# =========================================================
doc.add_page_break()
heading("6. Analysis", level=1)

heading("6.1 Judge Model Quality Is the Biggest Variable", level=2)
p("The 3B Ollama model underscored its own answers by over 30% on faithfulness. "
  "When the same answers were re-judged by Groq 70B, faithfulness jumped from 56.6% to 88.0% "
  "and context recall from 38.6% to 100%. The answers didn't change — only the judge improved. "
  "This confirms that small models are unreliable evaluators.")

heading("6.2 Generation Quality: 70B Is Modestly Better", level=2)
p("With the same Groq 70B judge, the actual generation gap is modest: "
  "12% on faithfulness (88% vs 100%) and 11.8% on BLEU (51.9% vs 63.7%). "
  "The 70B model generates more precise answers that stay closer to the source material.")

heading("6.3 Token-Overlap vs LLM-Judged Metrics", level=2)
p("Batch token-overlap faithfulness (76.6%) and LLM-judged faithfulness (88.0%) "
  "broadly agree but diverge on nuanced cases. Token-overlap is fast and free but misses "
  "semantic equivalence. LLM-judged metrics capture meaning but depend on judge quality. "
  "Both flagged the same weak spots (Type 1 Diabetes, MRSA).")

heading("6.4 QA-Pair Chunking Is Best for This Dataset", level=2)
p("The original QA-pair chunking outperformed both alternatives across all metrics. "
  "This is because the documents are already structured as question-answer pairs — "
  "chunking by QA boundaries preserves complete semantic units. "
  "Recursive splitting (84% recall) and semantic chunking (60% recall) both fragment "
  "these natural units, hurting retrieval quality.")

heading("6.5 More Chunks ≠ Better Retrieval", level=2)
p("Semantic chunking created 4,248 chunks (2x more than QA-pair) but scored lowest "
  "on every metric. Fragmentation increases noise in retrieval results.")

# =========================================================
# SECTION 7: KEY TAKEAWAYS
# =========================================================
heading("7. Key Takeaways", level=1)

p("Use a strong model (70B+) as the evaluation judge. Small models (3B) produce unreliable scores "
  "that can mislead evaluation by 30+ percentage points.", style_name="List Number")
p("Separate the judge from the generator to avoid self-evaluation bias. "
  "The fair comparison approach gives the most trustworthy results.", style_name="List Number")
p("BLEU scores are the most objective metric — use them alongside LLM-judged metrics "
  "to cross-validate results.", style_name="List Number")
p("Match chunking strategy to document structure. QA-pair chunking is ideal for structured "
  "FAQ-style documents. Use recursive splitting for unstructured prose.", style_name="List Number")
p("Token-overlap metrics are useful for fast iteration during development. "
  "Use LLM-judged metrics for thorough evaluation before deployment.", style_name="List Number")
p("The Groq free tier (llama-3.3-70b) provides excellent evaluation quality "
  "with generous daily quotas — a practical choice for projects without budget for paid APIs.", style_name="List Number")

# =========================================================
# SECTION 8: FILES REFERENCE
# =========================================================
doc.add_page_break()
heading("8. Project Files Reference", level=1)

heading("8.1 Evaluation Scripts", level=2)
table(
    ["File", "Purpose", "Generator", "Judge"],
    [
        ["evaluation.py", "Original baseline", "Ollama (3B)", "Ollama (3B)"],
        ["evaluation_groq.py", "API-based evaluation", "Groq (70B)", "Groq (70B)"],
        ["evaluation_fair_comparison.py", "Fair comparison (same judge)", "Saved answers", "Groq (70B)"],
        ["batch_evaluation.py", "Fast algorithmic eval", "Ollama (3B)", "Token-overlap"],
        ["evaluation_chunking_comparison.py", "Chunking method comparison", "Groq (70B)", "Groq (70B)"],
    ],
)

heading("8.2 Ingestion Scripts", level=2)
table(
    ["File", "Chunking Method", "DB Path"],
    [
        ["ingest.py", "QA-Pair (original)", "db/chroma_db"],
        ["ingest_recursive.py", "Recursive Character Splitting", "db/chroma_db_recursive"],
        ["ingest_semantic.py", "Semantic Chunking", "db/chroma_db_semantic"],
    ],
)

heading("8.3 Core Pipeline", level=2)
table(
    ["File", "Purpose"],
    [
        ["retrieval.py", "RAG retrieval + reranking + answer generation"],
        ["generation.py", "Interactive question-answering interface"],
        ["app.py", "Streamlit web application"],
        ["questions.txt", "5 test questions"],
        ["docs/", "423 HSE medical condition markdown files"],
    ],
)

heading("8.4 Results & Reports", level=2)
table(
    ["File", "Contents"],
    [
        ["ragas_results.csv", "Ollama evaluation results"],
        ["ragas_results_groq.csv", "Groq evaluation results"],
        ["ragas_details.json", "Ollama detailed results with contexts"],
        ["ragas_details_groq.json", "Groq detailed results with contexts"],
        ["ragas_results_fair_comparison.csv", "Fair comparison results"],
        ["batch_results.csv", "Batch token-overlap results"],
        ["chunking_comparison_results.json", "Chunking method comparison results"],
    ],
)

# =========================================================
# SAVE
# =========================================================
doc.save("Master_Evaluation_Report.docx")
print("Saved: Master_Evaluation_Report.docx")
