"""Generate 4 Word documents in ~/Downloads/ (outside project folder)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.expanduser("~/Downloads")

def new_doc():
    d = Document()
    s = d.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
    return d

def heading(d, t, level=1):
    h = d.add_heading(t, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

def tbl(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    d.add_paragraph()

def p(d, text, style=None):
    d.add_paragraph(text, style=style)


# =============================================================
# DOCUMENT 1: Pipeline Execution Order
# =============================================================
def doc1_execution_order():
    d = new_doc()
    t = d.add_heading("RAG Pipeline — File Execution Order", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    heading(d, "1. Data Collection & Preparation (Run Once)", level=1)

    heading(d, "Step 1: scrape_hse.py", level=2)
    p(d, "Run first. This script scrapes medical condition pages from the HSE (Health Service Executive) website. "
         "It collects condition names, sections (symptoms, causes, treatment, etc.), and their content. "
         "Output: hse_conditions_long.csv")
    p(d, "Command: python scrape_hse.py")

    heading(d, "Step 2: convert_hse_to_qa.py", level=2)
    p(d, "Run second. Takes the scraped CSV and converts each section into a question-answer pair. "
         "For example, a 'symptoms' section for 'Asthma' becomes: Q: 'What are the symptoms of asthma?' "
         "A: [the scraped content]. Output: hse_conditions_qa.csv")
    p(d, "Command: python convert_hse_to_qa.py")

    heading(d, "Step 3: generate_markdown.py", level=2)
    p(d, "Run third. Converts the QA CSV into structured markdown files — one per condition. "
         "Each file contains all QA pairs for that condition with metadata (section, source URL). "
         "Output: 423 .md files in the docs/ folder")
    p(d, "Command: python generate_markdown.py")

    d.add_page_break()
    heading(d, "2. Ingestion — Building the Vector Store", level=1)

    heading(d, "Step 4: ingest.py (Primary — QA-Pair Chunking)", level=2)
    p(d, "Run fourth. Parses the markdown files, splits them into QA-pair chunks, generates embeddings "
         "using sentence-transformers (all-MiniLM-L6-v2), and stores them in ChromaDB. "
         "This is the primary vector store used by the RAG pipeline. Output: db/chroma_db/")
    p(d, "Command: python ingest.py")

    heading(d, "Step 4a: ingest_recursive.py (Alternative Chunking)", level=2)
    p(d, "Optional. Uses RecursiveCharacterTextSplitter (800 chars, 200 overlap) instead of QA-pair chunking. "
         "Creates a separate vector store for comparison. Output: db/chroma_db_recursive/")
    p(d, "Command: python ingest_recursive.py")

    heading(d, "Step 4b: ingest_semantic.py (Alternative Chunking)", level=2)
    p(d, "Optional. Uses SemanticChunker that detects meaning shifts via embeddings. "
         "Creates a separate vector store for comparison. Output: db/chroma_db_semantic/")
    p(d, "Command: python ingest_semantic.py")

    d.add_page_break()
    heading(d, "3. Running the Application", level=1)

    heading(d, "Step 5: app.py (Streamlit Web Interface)", level=2)
    p(d, "The main user-facing application. Provides a web interface where users can select a disease, "
         "choose a question, or type their own query. Uses the retrieval pipeline to generate answers. "
         "Requires Ollama to be running locally.")
    p(d, "Command: streamlit run app.py")

    heading(d, "Supporting files used by app.py:", level=2)
    tbl(d, ["File", "Role"],
        [["retrieval.py", "Core RAG pipeline: retrieval, reranking, answer generation"],
         ["generation.py", "Interactive CLI interface for question-answering (alternative to app.py)"]])

    d.add_page_break()
    heading(d, "4. Evaluation (Run After Pipeline is Working)", level=1)

    tbl(d, ["Order", "File", "What It Does", "Prerequisites"],
        [["1", "evaluation.py", "Baseline evaluation using Ollama\nas both generator and judge", "Ollama running + db/chroma_db"],
         ["2", "evaluation_groq.py", "Same evaluation using Groq API\nas both generator and judge", "GROQ_API_KEY in .env"],
         ["3", "evaluation_fair_comparison.py", "Re-judges both answer sets\nwith same Groq 70B judge", "Results from steps 1 & 2"],
         ["4", "batch_evaluation.py", "Fast token-overlap evaluation\n(no LLM judge needed)", "Ollama running + db/chroma_db"],
         ["5", "evaluation_chunking_comparison.py", "Compares 3 chunking methods", "All 3 vector stores built"],
         ["6", "evaluation_improvements.py\n+ run_improvements.py\n+ run_single_experiment.py", "Compares FAISS, MPNet,\nand hybrid search", "GROQ_API_KEY + db/chroma_db"]])

    d.add_page_break()
    heading(d, "5. Report Generation", level=1)
    p(d, "These scripts generate Word documents with charts and analysis. "
         "Run them after the corresponding evaluations are complete.")
    tbl(d, ["File", "Generates"],
        [["generate_comparison.py", "Original Ollama vs Groq comparison charts"],
         ["generate_word_report.py", "First comparison Word report"],
         ["generate_full_report.py", "Full comparison report (original + fair)"],
         ["generate_batch_vs_fair_report.py", "Batch vs Fair comparison report"],
         ["generate_chunking_report.py", "Chunking method comparison report"],
         ["generate_improvements_report.py", "RAG improvements report"],
         ["generate_master_report.py", "Master report combining all evaluations"]])

    d.add_page_break()
    heading(d, "6. Complete Execution Flow Diagram", level=1)
    p(d, "scrape_hse.py → convert_hse_to_qa.py → generate_markdown.py → ingest.py → app.py")
    p(d, "")
    p(d, "For evaluation:")
    p(d, "evaluation.py → evaluation_groq.py → evaluation_fair_comparison.py")
    p(d, "batch_evaluation.py (independent)")
    p(d, "ingest_recursive.py + ingest_semantic.py → evaluation_chunking_comparison.py")
    p(d, "run_improvements.py (runs run_single_experiment.py for each experiment)")

    d.save(os.path.join(OUT, "1_Pipeline_Execution_Order.docx"))
    print("Saved: 1_Pipeline_Execution_Order.docx")


# =============================================================
# DOCUMENT 2: Evaluations & Comparisons Explained for Audience
# =============================================================
def doc2_evaluations_explained():
    d = new_doc()
    t = d.add_heading("RAG Evaluation & Comparison — Audience Guide", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    heading(d, "1. What We Are Evaluating", level=1)
    p(d, "We built a Retrieval-Augmented Generation (RAG) pipeline that answers medical questions "
         "by retrieving relevant information from 423 HSE medical condition documents and generating "
         "answers using a Large Language Model (LLM). The evaluation measures how good these answers are.")

    heading(d, "1.1 The Three Core Metrics", level=2)
    tbl(d, ["Metric", "What It Measures", "How It Works", "Why It Matters"],
        [["Faithfulness", "Is the answer grounded\nin the retrieved context?", "An LLM judge reads the context\nand answer, scores 0.0–1.0", "Prevents hallucination —\nthe answer should only contain\nwhat the documents say"],
         ["Context Recall", "Did we retrieve the\nright documents?", "An LLM judge compares retrieved\ncontext to the ground truth", "If we retrieve wrong documents,\neven a perfect LLM can't\ngenerate a good answer"],
         ["BLEU Score", "Does the answer match\nthe expected answer?", "Algorithmic word overlap\nbetween answer and reference", "Objective measure — no LLM\nbias, purely mathematical"]])

    d.add_page_break()
    heading(d, "2. Evaluation 1: Ollama vs Groq (Original)", level=1)
    p(d, "The first comparison tested two LLM backends: a small local model (Ollama llama3.2, 3 billion parameters) "
         "versus a large cloud model (Groq llama-3.3-70b, 70 billion parameters).")
    p(d, "Key finding: Groq scored dramatically higher — 100% faithfulness vs 56.6% for Ollama. "
         "But this was misleading because each model judged its own answers. The small model was a poor judge.")

    tbl(d, ["Metric", "Ollama (3B)", "Groq (70B)"],
        [["Faithfulness", "56.6%", "100.0%"],
         ["Context Recall", "38.6%", "100.0%"],
         ["BLEU Score", "51.9%", "63.7%"]])

    p(d, "How to explain to audience: \"The initial results showed a massive gap, but we discovered "
         "this was mostly because the small model couldn't evaluate properly — like having a student "
         "grade their own exam versus having a professor grade it.\"")

    d.add_page_break()
    heading(d, "3. Evaluation 2: Fair Comparison (Same Judge)", level=1)
    p(d, "To get a fair comparison, I re-evaluated both sets of answers using the same judge: "
         "Groq's 70B model. This isolated the actual answer quality from the judge quality.")

    tbl(d, ["Metric", "Ollama Answers\n(Groq-judged)", "Groq Answers\n(Groq-judged)"],
        [["Faithfulness", "88.0%", "100.0%"],
         ["Context Recall", "100.0%", "100.0%"],
         ["BLEU Score", "51.9%", "63.7%"]])

    p(d, "How to explain: \"When we used the same judge for both, Ollama's faithfulness jumped from "
         "56.6% to 88.0% — the answers were always decent, the judge was the problem. "
         "The real quality gap is only about 12%, not 43%.\"")

    d.add_page_break()
    heading(d, "4. Evaluation 3: Batch Evaluation (Token-Overlap)", level=1)
    p(d, "I also implemented a fast, algorithmic evaluation that doesn't need any LLM. "
         "It uses simple token overlap to measure faithfulness, coverage, and relevance.")

    tbl(d, ["Metric", "Score", "How It Works"],
        [["Faithfulness", "76.6%", "What % of answer words appear in the context"],
         ["Coverage", "29.4%", "What % of context words appear in the answer"],
         ["Relevance", "57.6%", "What % of question words appear in the answer"]])

    p(d, "How to explain: \"This gives us a quick sanity check without needing an API. "
         "It's like spell-checking versus having someone proofread — fast but shallow. "
         "Both methods flagged the same weak spots, which validates our results.\"")

    d.add_page_break()
    heading(d, "5. Evaluation 4: Chunking Method Comparison", level=1)
    p(d, "I tested three different ways of splitting documents into chunks for the vector store:")

    tbl(d, ["Method", "How It Works", "Faithfulness", "Context Recall", "BLEU"],
        [["QA-Pair (Original)", "Split by question-answer boundaries", "90.0%", "100.0%", "75.4%"],
         ["Recursive Character", "Split at 800 chars with 200 overlap", "76.0%", "84.0%", "58.8%"],
         ["Semantic Chunking", "Split where meaning shifts (via embeddings)", "76.0%", "60.0%", "55.9%"]])

    p(d, "How to explain: \"Since our documents are already structured as Q&A pairs, "
         "chunking by those natural boundaries works best. Generic splitting methods break "
         "questions apart from their answers, hurting retrieval. This shows that chunking strategy "
         "should match your document structure.\"")

    d.add_page_break()
    heading(d, "6. Evaluation 5: RAG Improvements", level=1)
    p(d, "Finally, I tested specific technical improvements to the pipeline:")

    tbl(d, ["Improvement", "What Changed", "Faithfulness", "BLEU", "Speed"],
        [["Baseline", "ChromaDB + MiniLM embedding", "76.0%", "77.4%", "4,978ms"],
         ["FAISS", "Replaced ChromaDB with FAISS", "84.0%", "77.4%", "1,453ms"],
         ["Better Embedding", "Used MPNet (768 dim) instead of MiniLM (384 dim)", "84.0%", "80.8%", "1,409ms"],
         ["Hybrid Search", "Combined BM25 keyword + vector search", "88.0%", "81.0%", "1,313ms"]])

    p(d, "How to explain: \"Each improvement targets a different bottleneck. FAISS made retrieval 3.4x faster. "
         "MPNet embeddings improved answer quality by capturing more meaning. Hybrid search gave the best "
         "overall results by combining keyword matching with semantic search — catching documents that "
         "either method alone would miss.\"")

    d.add_page_break()
    heading(d, "7. Summary: What I Chose and Why", level=1)
    tbl(d, ["Component", "My Choice", "Why"],
        [["Chunking", "QA-Pair", "Documents are structured as Q&A — matches naturally"],
         ["Vector Store", "FAISS", "3.4x faster than ChromaDB, same quality"],
         ["Embedding", "all-mpnet-base-v2", "Richer 768-dim embeddings, better retrieval"],
         ["Retrieval", "Hybrid (BM25 + Vector)", "Best quality — covers both keyword and semantic matching"],
         ["LLM (Generation)", "Groq llama-3.3-70b", "Best answer quality, fast API, free tier"],
         ["LLM (Evaluation)", "Groq llama-3.3-70b (separate judge)", "Reliable evaluation, avoids self-judging bias"]])

    p(d, "How to explain: \"I systematically tested each component of the RAG pipeline — the chunking, "
         "the vector store, the embeddings, and the retrieval strategy. Each experiment isolated one variable "
         "so I could measure its impact. The final configuration combines the best choice for each component.\"")

    d.save(os.path.join(OUT, "2_Evaluations_Explained_For_Audience.docx"))
    print("Saved: 2_Evaluations_Explained_For_Audience.docx")


# =============================================================
# DOCUMENT 3: Explanation of Each Python File
# =============================================================
def doc3_file_explanations():
    d = new_doc()
    t = d.add_heading("Python File Explanations", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph()

    files = [
        ("scrape_hse.py", "Data Collection — Web Scraper",
         "This script scrapes medical condition pages from the HSE (Health Service Executive) website at https://www2.hse.ie/conditions/. "
         "It collects all condition links from the index page, then visits each condition page to extract structured content.",
         "Key functions:\n"
         "- extract_condition_links(): Finds all condition URLs from the HSE index page\n"
         "- scrape_condition(url): Extracts condition name, sections (symptoms, causes, treatment), and content from each page\n"
         "- normalize_heading(): Maps varied heading text to standardised section names (e.g., 'Symptoms of X' → 'symptoms')\n"
         "- extract_section_content(): Parses HTML to extract text, lists, and callout boxes from each section\n\n"
         "Output: hse_conditions_long.csv with columns: condition_name, section, content, source_url, source_path\n"
         "The script includes a 1-second delay between requests to be respectful to the HSE server."),

        ("convert_hse_to_qa.py", "Data Transformation — CSV to QA Pairs",
         "Converts the raw scraped data into question-answer format. Each section of a condition becomes a QA pair "
         "using predefined question templates.",
         "Key logic:\n"
         "- SECTION_QUESTION_MAP: Maps section types to question templates. For example, 'symptoms' → 'What are the symptoms of {condition}?'\n"
         "- For sections without a template, generates: 'What should I know about {condition} and {section}?'\n"
         "- Filters out junk sections (contents, support links) and answers shorter than 30 characters\n"
         "- Deduplicates by condition + section + answer\n\n"
         "Output: hse_conditions_qa.csv with columns: chunk_id, condition_name, question, answer, section, source_url"),

        ("generate_markdown.py", "Data Formatting — QA CSV to Markdown Files",
         "Converts the QA CSV into structured markdown files, one per medical condition. "
         "These markdown files become the source documents for the RAG pipeline.",
         "Key logic:\n"
         "- Groups QA pairs by condition_name\n"
         "- Sorts sections in a logical order (overview → symptoms → tests → treatment → causes → prevention)\n"
         "- Generates markdown with ## Question / ### Answer / **Section:** / **Source:** structure\n"
         "- Creates clean filenames from condition names (e.g., 'Type 1 diabetes' → 'type_1_diabetes.md')\n\n"
         "Output: 423 .md files in the docs/ folder"),

        ("ingest.py", "Ingestion — QA-Pair Chunking (Primary)",
         "The primary ingestion script. Parses markdown files into QA-pair chunks, generates embeddings, "
         "and stores them in ChromaDB.",
         "Key functions:\n"
         "- split_into_question_blocks(): Splits markdown by '## Question' headers\n"
         "- parse_question_block(): Extracts question, answer, section, and source URL from each block\n"
         "- split_long_answer(): Splits answers longer than 800 characters at paragraph/sentence boundaries\n"
         "- make_question_hash(): Creates a unique hash for each QA pair for deduplication\n"
         "- build_chunk_text(): Reconstructs the chunk with disease title, question, answer, and metadata\n\n"
         "Uses sentence-transformers/all-MiniLM-L6-v2 for embeddings (384 dimensions)\n"
         "Output: db/chroma_db/ (ChromaDB vector store with ~1,950 chunks)"),

        ("ingest_recursive.py", "Ingestion — Recursive Character Splitting (Alternative)",
         "Alternative ingestion using LangChain's RecursiveCharacterTextSplitter. "
         "Splits text at natural boundaries (paragraphs → sentences → words) with overlap.",
         "Configuration:\n"
         "- Chunk size: 800 characters\n"
         "- Overlap: 200 characters\n"
         "- Separators: paragraph breaks, line breaks, sentences, spaces\n\n"
         "This method doesn't understand the QA structure — it just splits by character count. "
         "Created for comparison with QA-pair chunking.\n"
         "Output: db/chroma_db_recursive/ (2,801 chunks)"),

        ("ingest_semantic.py", "Ingestion — Semantic Chunking (Alternative)",
         "Alternative ingestion using LangChain's SemanticChunker. "
         "Uses the embedding model to detect where the meaning shifts in the text and splits at those boundaries.",
         "Configuration:\n"
         "- Breakpoint type: percentile\n"
         "- Threshold: 70 (splits when similarity drops below 70th percentile)\n\n"
         "Groups semantically similar sentences together. Works well for free-form prose but "
         "poorly for structured markdown where metadata lines create artificial boundaries.\n"
         "Output: db/chroma_db_semantic/ (4,248 chunks)"),

        ("retrieval.py", "Core RAG Pipeline — Retrieval & Generation",
         "The heart of the RAG system. Handles query processing, document retrieval, reranking, and answer generation.",
         "Key functions:\n"
         "- extract_disease_from_query(): Identifies which disease the query is about by scanning the vector store\n"
         "- detect_section(): Identifies the query type (symptoms, causes, treatment) from keywords\n"
         "- similarity_retrieve() / mmr_retrieve(): Two retrieval strategies — pure similarity and Maximum Marginal Relevance\n"
         "- apply_metadata_boosts(): Boosts documents matching the detected disease/section\n"
         "- rerank_documents(): Uses a cross-encoder (ms-marco-MiniLM-L6-v2) to rerank candidates\n"
         "- select_top_question_groups(): Groups chunks by question hash and merges multi-part answers\n"
         "- ask(): The main entry point — takes a query, returns answer + sources + metadata\n\n"
         "Uses Ollama (llama3.2) for generation by default."),

        ("generation.py", "Interactive CLI Interface",
         "Provides a command-line interface for asking questions interactively. "
         "An alternative to the Streamlit web app.",
         "Key functions:\n"
         "- format_sources(): Formats source information for display\n"
         "- save_history(): Saves each Q&A interaction to generation_history.json\n"
         "- display_result(): Prints the answer, retrieval info, and sources\n\n"
         "Uses retrieval.py's ask() function internally."),

        ("app.py", "Streamlit Web Application",
         "The main user-facing web interface built with Streamlit. Provides disease selection, "
         "question browsing, custom queries, and debug information.",
         "Key features:\n"
         "- Disease dropdown populated from the vector store metadata\n"
         "- Question suggestions for each disease\n"
         "- Custom query input\n"
         "- Configurable retrieval method (similarity, MMR, hybrid)\n"
         "- Confidence scoring based on reranker scores\n"
         "- Debug mode showing retrieved chunks and scores\n\n"
         "Command: streamlit run app.py"),
    ]

    for fname, title, desc, details in files:
        heading(d, f"{fname} — {title}", level=1)
        p(d, desc)
        p(d, details)
        d.add_paragraph()

    d.add_page_break()
    # Evaluation files
    heading(d, "Evaluation Files", level=0)

    eval_files = [
        ("evaluation.py", "Baseline Evaluation (Ollama)",
         "The original evaluation script. Uses Ollama (llama3.2, 3B) as both the answer generator and the evaluation judge.",
         "Key functions:\n"
         "- compute_bleu(): Calculates unigram BLEU score between reference and candidate answers\n"
         "- evaluate_faithfulness(): Prompts the LLM to rate if the answer is grounded in context (0.0–1.0)\n"
         "- evaluate_context_recall(): Prompts the LLM to rate if context covers the reference answer\n"
         "- build_reference_map(): Extracts ground-truth QA pairs from markdown files\n"
         "- run_rag_on_questions(): Runs the RAG pipeline on each question and collects results\n"
         "- run_evaluation(): Scores all results using the LLM judge + BLEU\n\n"
         "Output: ragas_results.csv, ragas_details.json"),

        ("evaluation_groq.py", "API-Based Evaluation (Groq)",
         "Same evaluation logic as evaluation.py but uses Groq's llama-3.3-70b-versatile (70B) via API "
         "instead of local Ollama. Includes retry logic with exponential backoff for rate limiting.",
         "Key additions over evaluation.py:\n"
         "- _make_groq_llm(): Creates a ChatGroq instance using GROQ_API_KEY from .env\n"
         "- invoke_with_retry(): Handles API rate limits with exponential backoff\n"
         "- _RetryLLMWrapper: Wraps the LLM so retrieval.py's ask() also uses Groq with retry\n"
         "- Patches retrieval.load_llm to swap Ollama for Groq in the RAG pipeline\n\n"
         "Output: ragas_results_groq.csv, ragas_details_groq.json"),

        ("evaluation_fair_comparison.py", "Fair Comparison (Same Judge)",
         "Re-evaluates both Ollama and Groq generated answers using the same Groq 70B judge. "
         "Loads existing answers from ragas_details.json and ragas_details_groq.json — does not regenerate them.",
         "Purpose: Isolates answer quality from judge quality. Shows that the original gap was mostly "
         "due to the weak 3B judge, not poor answer quality.\n\n"
         "Output: ragas_results_fair_comparison.csv, ragas_details_fair_comparison.json, Fair_Comparison_Report.docx"),

        ("batch_evaluation.py", "Fast Algorithmic Evaluation (Token-Overlap)",
         "Uses simple token-overlap metrics instead of an LLM judge. Fast, deterministic, and free.",
         "Metrics:\n"
         "- context_overlap (faithfulness): answer tokens ∩ context tokens / answer tokens\n"
         "- context_coverage: answer tokens ∩ context tokens / context tokens\n"
         "- relevance_score: query tokens ∩ answer tokens / query tokens\n\n"
         "No LLM is involved in scoring — purely algorithmic.\n"
         "Output: batch_results.csv, batch_details.json"),

        ("evaluation_chunking_comparison.py", "Chunking Method Comparison",
         "Evaluates all three chunking methods (QA-pair, recursive, semantic) using the same Groq 70B "
         "generator and judge. Each method uses its own vector store.",
         "Key design:\n"
         "- Standalone retrieval pipeline (doesn't modify retrieval.py)\n"
         "- Loads each vector store independently\n"
         "- Uses cross-encoder reranking for all methods\n"
         "- Same questions, same judge, only chunking differs\n\n"
         "Output: chunking_comparison_results.json"),

        ("run_improvements.py + run_single_experiment.py", "RAG Improvement Experiments",
         "Tests four pipeline configurations: Baseline (ChromaDB+MiniLM), FAISS, FAISS+MPNet, and Hybrid Search. "
         "Each experiment runs in a separate subprocess to avoid memory conflicts.",
         "run_improvements.py: Orchestrator that launches each experiment as a subprocess\n"
         "run_single_experiment.py: Runs a single experiment (baseline/faiss/mpnet/hybrid)\n\n"
         "Key improvements tested:\n"
         "- FAISS: Faster vector search using Facebook's optimised library\n"
         "- MPNet: Higher quality embeddings (768 dim vs 384 dim)\n"
         "- Hybrid: BM25 keyword search + vector similarity, merged and reranked\n\n"
         "Output: improvement_*.json, improvement_results.json"),
    ]

    for fname, title, desc, details in eval_files:
        heading(d, f"{fname} — {title}", level=1)
        p(d, desc)
        p(d, details)
        d.add_paragraph()

    d.save(os.path.join(OUT, "3_Python_File_Explanations.docx"))
    print("Saved: 3_Python_File_Explanations.docx")


# =============================================================
# DOCUMENT 4: Evaluation Challenges & How I Overcame Them
# =============================================================
def doc4_challenges_story():
    d = new_doc()
    t = d.add_heading("Evaluation Challenges & How I Overcame Them", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A journey through building and evaluating a Medical RAG Pipeline")
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55); r.italic = True
    d.add_page_break()

    heading(d, "Challenge 1: Getting the Evaluation Pipeline Running", level=1)
    p(d, "When I first tried to run evaluation.py, it wouldn't start. The script depended on Ollama "
         "running as a local server, but the Ollama service wasn't active on my machine. "
         "I got connection errors every time the script tried to call the LLM.")
    heading(d, "How I Solved It", level=2)
    p(d, "I realised Ollama needs to be running as a background service before any script can use it. "
         "I started the Ollama server with 'ollama serve' and verified the llama3.2 model was available "
         "with 'ollama list'. Once the server was running, evaluation.py executed successfully and "
         "produced the first set of results: Faithfulness 56.6%, Context Recall 38.6%, BLEU 51.9%.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Local LLM setups require the inference server to be running. This is a common gotcha "
         "that's easy to miss, especially when switching between local and API-based models.")

    d.add_page_break()
    heading(d, "Challenge 2: Gemini API Rate Limits", level=1)
    p(d, "I wanted to compare the local Ollama results with a cloud-based LLM. I created evaluation_gemini.py "
         "using Google's Gemini API with my free-tier API key. The script started fine and processed the first "
         "question, but then crashed with a 429 RESOURCE_EXHAUSTED error.")
    p(d, "The error message said I had exceeded my quota for gemini-2.0-flash — both the per-minute "
         "and daily request limits were at zero. My free-tier key had been exhausted.")
    heading(d, "What I Tried", level=2)
    p(d, "First, I obtained a new Gemini API key, hoping the quota would reset. But the new key "
         "hit the same rate limits. I added retry logic with exponential backoff (waiting 30s, 60s, 120s, 240s "
         "between retries), but the daily quota was exhausted — no amount of waiting within the same day would help.")
    heading(d, "How I Solved It", level=2)
    p(d, "I researched alternative free-tier LLM APIs and found Groq, which offers a generous free tier "
         "with around 14,400 requests per day. I installed langchain-groq, created evaluation_groq.py using "
         "Groq's llama-3.3-70b-versatile model, and it ran perfectly on the first attempt — all 15 LLM calls "
         "completed without any rate limiting.")
    p(d, "I removed the Gemini evaluation file since it never produced results, keeping the project clean.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Free-tier API keys have strict quotas that can be exhausted quickly during evaluation "
         "(which makes many LLM calls). Always check the rate limits before committing to a provider. "
         "Groq's free tier turned out to be ideal for evaluation workloads.")

    d.add_page_break()
    heading(d, "Challenge 3: Misleading Evaluation Results", level=1)
    p(d, "When I compared the Ollama and Groq results, the gap was enormous: Ollama scored 56.6% "
         "faithfulness while Groq scored 100%. At first, I thought the 70B model was simply that much better.")
    p(d, "But something didn't add up. The BLEU scores (which are purely algorithmic) only showed "
         "a modest improvement: 51.9% → 63.7%. If the answers were truly that different, "
         "BLEU should have shown a bigger gap too.")
    heading(d, "How I Investigated", level=2)
    p(d, "I realised the problem: each model was judging its own answers. The 3B model was both "
         "generating answers AND evaluating them. A small model is a weak evaluator — it can't "
         "reliably assess whether an answer is faithful to the context.")
    heading(d, "How I Solved It", level=2)
    p(d, "I created evaluation_fair_comparison.py that loads the saved answers from both models "
         "and re-evaluates ALL of them using the same Groq 70B judge. This isolated the variable: "
         "same judge, different answers.")
    p(d, "The results were revealing: Ollama's faithfulness jumped from 56.6% to 88.0% — "
         "the answers were always decent, the 3B judge just couldn't evaluate them properly. "
         "The real quality gap was only about 12%, not 43%.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Never use the same model as both generator and judge, especially a small model. "
         "Always use a strong, separate model for evaluation. BLEU scores serve as a useful "
         "sanity check because they have no LLM bias.")

    d.add_page_break()
    heading(d, "Challenge 4: Choosing the Right Chunking Strategy", level=1)
    p(d, "I knew chunking affects retrieval quality, but I wasn't sure which method would work best "
         "for my structured medical documents. I needed to test multiple approaches.")
    heading(d, "What I Tried", level=2)
    p(d, "I implemented three chunking methods:\n"
         "1. QA-Pair chunking (the original) — splits by question-answer boundaries\n"
         "2. Recursive Character splitting — splits at 800 characters with 200 overlap\n"
         "3. Semantic chunking — uses embeddings to detect meaning shifts")
    p(d, "Each method required its own ingestion script and vector store. I had to ensure "
         "they didn't interfere with each other or the original database.")
    heading(d, "How I Solved It", level=2)
    p(d, "I created separate ingestion scripts (ingest_recursive.py, ingest_semantic.py) that write "
         "to their own database directories (db/chroma_db_recursive, db/chroma_db_semantic). "
         "Then I built evaluation_chunking_comparison.py with its own standalone retrieval pipeline "
         "that could point to any vector store without modifying retrieval.py.")
    p(d, "The results clearly showed QA-pair chunking winning across all metrics (90% faithfulness, "
         "100% context recall, 75.4% BLEU). This made sense because the documents are already "
         "structured as Q&A pairs — splitting by those natural boundaries preserves semantic completeness.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Chunking strategy should match document structure. More chunks doesn't mean better retrieval — "
         "semantic chunking created 2x more chunks but scored lowest. Always test multiple approaches "
         "and let the data decide.")

    d.add_page_break()
    heading(d, "Challenge 5: Memory Crashes with FAISS", level=1)
    p(d, "When I tried to run all four improvement experiments (Baseline, FAISS, MPNet, Hybrid) "
         "in a single Python process, the script crashed with a segmentation fault. "
         "Loading ChromaDB, FAISS, two different embedding models, and the cross-encoder reranker "
         "simultaneously exceeded the available memory.")
    heading(d, "How I Solved It", level=2)
    p(d, "I restructured the experiment runner into two files:\n"
         "- run_improvements.py: An orchestrator that launches each experiment as a separate subprocess\n"
         "- run_single_experiment.py: Runs one experiment at a time, then exits\n\n"
         "Each experiment runs in its own Python process with fresh memory. The orchestrator collects "
         "results from JSON files written by each subprocess. This solved the memory issue completely — "
         "all four experiments ran successfully.")
    heading(d, "Lesson Learned", level=2)
    p(d, "When working with multiple ML models and vector stores, memory management matters. "
         "Running experiments in separate processes is a clean solution that avoids conflicts "
         "between different libraries (FAISS, ChromaDB, multiple embedding models).")

    d.add_page_break()
    heading(d, "Challenge 6: Ensuring Fair Comparisons", level=1)
    p(d, "Throughout the evaluation process, I had to be careful that each comparison was fair — "
         "changing only one variable at a time while keeping everything else constant.")
    heading(d, "How I Managed It", level=2)
    p(d, "I followed a strict methodology:\n"
         "- For model comparison: Same questions, same retrieval, same judge — only the generator changed\n"
         "- For chunking comparison: Same questions, same LLM, same judge — only the chunking changed\n"
         "- For improvement experiments: Same questions, same judge — only the vector store/embedding/retrieval changed\n"
         "- Every new experiment created separate output files and databases, never overwriting previous results")
    p(d, "I also kept the original evaluation.py completely untouched throughout the entire process. "
         "Every new experiment was a new file. This meant I could always go back and verify earlier results.")
    heading(d, "Lesson Learned", level=2)
    p(d, "Scientific evaluation requires controlled experiments. Change one variable at a time, "
         "keep everything else constant, and never overwrite previous results. "
         "This discipline made it possible to draw clear conclusions from each comparison.")

    d.add_page_break()
    heading(d, "Summary: The Journey", level=1)
    p(d, "Building this RAG pipeline was straightforward. Evaluating it properly was the real challenge. "
         "Along the way, I learned that:")
    p(d, "1. Local LLM setup requires attention to service management (Ollama server)", style="List Number")
    p(d, "2. Free API tiers have real limits — always have a backup plan (Gemini → Groq)", style="List Number")
    p(d, "3. Self-evaluation is unreliable — always use a separate, strong judge model", style="List Number")
    p(d, "4. Chunking strategy should match document structure, not be chosen arbitrarily", style="List Number")
    p(d, "5. Memory management matters when combining multiple ML models", style="List Number")
    p(d, "6. Fair comparison requires strict experimental discipline", style="List Number")
    p(d, "")
    p(d, "Each challenge taught me something about building production-quality AI systems. "
         "The final pipeline — with QA-pair chunking, FAISS, MPNet embeddings, hybrid search, "
         "and Groq's 70B model — represents the best configuration I found through systematic experimentation.")

    d.save(os.path.join(OUT, "4_Evaluation_Challenges_Story.docx"))
    print("Saved: 4_Evaluation_Challenges_Story.docx")


# =============================================================
# RUN ALL
# =============================================================
if __name__ == "__main__":
    doc1_execution_order()
    doc2_evaluations_explained()
    doc3_file_explanations()
    doc4_challenges_story()
    print(f"\nAll 4 documents saved to: {OUT}")
