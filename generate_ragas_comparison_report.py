"""Generate RAGAS vs Custom comparison chart and Word report."""
import json
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

with open("ragas_vs_custom_details.json") as f:
    data = json.load(f)["comparison"]

QS = ["Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention", "Type 1\nDiabetes", "MRSA"]
QF = [d["question"] for d in data]
C_RAGAS = "#E8833A"
C_CUSTOM = "#4A90D9"
C_PREC = "#2ECC71"
avg = lambda lst, k: sum(r[k] for r in lst) / len(lst)

# =========================================================
# CHARTS
# =========================================================
# Chart 1: Overall comparison
fig, ax = plt.subplots(figsize=(9, 5.5))
metrics = ["Faithfulness", "Context Recall"]
x = np.arange(len(metrics))
w = 0.3
ragas_vals = [avg(data, "ragas_faithfulness")*100, avg(data, "ragas_context_recall")*100]
custom_vals = [avg(data, "custom_faithfulness")*100, avg(data, "custom_context_recall")*100]
b1 = ax.bar(x - w/2, ragas_vals, w, label="RAGAS (Groq 70B judge)", color=C_RAGAS)
b2 = ax.bar(x + w/2, custom_vals, w, label="Custom (Ollama 3B judge)", color=C_CUSTOM)
for b in b1:
    ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=9)
for b in b2:
    ax.text(b.get_x()+b.get_width()/2, b.get_height()+1.5, f"{b.get_height():.1f}%", ha="center", fontsize=9)
ax.set_ylabel("Score (%)"); ax.set_title("RAGAS vs Custom Evaluation: Overall Averages", fontsize=13, fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(metrics); ax.set_ylim(0, 118)
ax.legend(fontsize=10); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("ragas_cmp_overall.png", dpi=150); plt.close()

# Chart 2: Faithfulness per question
fig, ax = plt.subplots(figsize=(11, 5.5))
x = np.arange(5); w = 0.32
ax.bar(x - w/2, [d["ragas_faithfulness"]*100 for d in data], w, label="RAGAS (Groq 70B)", color=C_RAGAS)
ax.bar(x + w/2, [d["custom_faithfulness"]*100 for d in data], w, label="Custom (Ollama 3B)", color=C_CUSTOM)
ax.set_ylabel("Score (%)"); ax.set_title("Faithfulness per Question: RAGAS vs Custom", fontsize=13, fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(QS, fontsize=9); ax.set_ylim(0, 118)
ax.legend(fontsize=10); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("ragas_cmp_faith.png", dpi=150); plt.close()

# Chart 3: Context Recall per question
fig, ax = plt.subplots(figsize=(11, 5.5))
ax.bar(x - w/2, [d["ragas_context_recall"]*100 for d in data], w, label="RAGAS (Groq 70B)", color=C_RAGAS)
ax.bar(x + w/2, [d["custom_context_recall"]*100 for d in data], w, label="Custom (Ollama 3B)", color=C_CUSTOM)
ax.set_ylabel("Score (%)"); ax.set_title("Context Recall per Question: RAGAS vs Custom", fontsize=13, fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(QS, fontsize=9); ax.set_ylim(0, 118)
ax.legend(fontsize=10); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("ragas_cmp_recall.png", dpi=150); plt.close()

# Chart 4: All RAGAS metrics (including Context Precision)
fig, ax = plt.subplots(figsize=(11, 5.5))
w2 = 0.25
ax.bar(x - w2, [d["ragas_faithfulness"]*100 for d in data], w2, label="Faithfulness", color=C_RAGAS)
ax.bar(x, [d["ragas_context_recall"]*100 for d in data], w2, label="Context Recall", color=C_CUSTOM)
ax.bar(x + w2, [d["ragas_context_precision"]*100 for d in data], w2, label="Context Precision", color=C_PREC)
ax.set_ylabel("Score (%)"); ax.set_title("RAGAS Evaluation: All Metrics per Question", fontsize=13, fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(QS, fontsize=9); ax.set_ylim(0, 118)
ax.legend(fontsize=10); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("ragas_cmp_all_ragas.png", dpi=150); plt.close()

print("Charts generated.")

# =========================================================
# WORD DOCUMENT
# =========================================================
doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)

def heading(t, level=1):
    h = doc.add_heading(t, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

def img(path):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# Title
doc.add_paragraph()
t = doc.add_heading("RAGAS vs Custom Evaluation Comparison", level=0)
for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Industry-Standard RAGAS Framework vs Custom Prompt-Based Evaluation")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
doc.add_page_break()

# Section 1: Overview
heading("1. Overview")
doc.add_paragraph(
    "This report compares two evaluation approaches applied to the same set of Ollama-generated answers. "
    "The answers were generated once using the local Ollama llama3.2 (3B) model and saved. "
    "Both evaluation methods then scored these identical answers.")

tbl(["", "Custom Evaluation", "RAGAS Library"],
    [["File", "evaluation.py", "evaluation_ragas_compare.py"],
     ["Judge Model", "Ollama llama3.2 (3B, local)", "Groq llama-3.3-70b (70B, API)"],
     ["Method", "Single prompt: 'Rate 0-1'", "Claim decomposition + verification"],
     ["Faithfulness", "LLM gives one holistic score", "Extracts claims, verifies each against context"],
     ["Context Recall", "LLM gives one holistic score", "Checks each reference sentence against context"],
     ["Context Precision", "Not available", "Checks if each context chunk was useful"],
     ["BLEU Score", "Algorithmic word overlap", "Not included"]])

# Section 2: Overall Results
doc.add_page_break()
heading("2. Overall Results")

tbl(["Metric", "RAGAS\n(Groq 70B judge)", "Custom\n(Ollama 3B judge)", "Difference"],
    [["Faithfulness", f"{avg(data,'ragas_faithfulness')*100:.1f}%", f"{avg(data,'custom_faithfulness')*100:.1f}%",
      f"{(avg(data,'ragas_faithfulness')-avg(data,'custom_faithfulness'))*100:+.1f}%"],
     ["Context Recall", f"{avg(data,'ragas_context_recall')*100:.1f}%", f"{avg(data,'custom_context_recall')*100:.1f}%",
      f"{(avg(data,'ragas_context_recall')-avg(data,'custom_context_recall'))*100:+.1f}%"],
     ["Context Precision", f"{avg(data,'ragas_context_precision')*100:.1f}%", "N/A", "—"],
     ["BLEU Score", "N/A", f"{avg(data,'custom_bleu')*100:.1f}%", "—"]])

img("ragas_cmp_overall.png")

# Section 3: Per-Question
doc.add_page_break()
heading("3. Per-Question Breakdown")

heading("3.1 Faithfulness", level=2)
tbl(["Question", "RAGAS", "Custom", "Diff"],
    [[QF[i][:45], f"{data[i]['ragas_faithfulness']*100:.0f}%", f"{data[i]['custom_faithfulness']*100:.0f}%",
      f"{(data[i]['ragas_faithfulness']-data[i]['custom_faithfulness'])*100:+.0f}%"] for i in range(5)])
img("ragas_cmp_faith.png")

heading("3.2 Context Recall", level=2)
tbl(["Question", "RAGAS", "Custom", "Diff"],
    [[QF[i][:45], f"{data[i]['ragas_context_recall']*100:.0f}%", f"{data[i]['custom_context_recall']*100:.0f}%",
      f"{(data[i]['ragas_context_recall']-data[i]['custom_context_recall'])*100:+.0f}%"] for i in range(5)])
img("ragas_cmp_recall.png")

heading("3.3 RAGAS: All Three Metrics", level=2)
tbl(["Question", "Faithfulness", "Context Recall", "Context Precision"],
    [[QF[i][:45], f"{data[i]['ragas_faithfulness']*100:.0f}%", f"{data[i]['ragas_context_recall']*100:.0f}%",
      f"{data[i]['ragas_context_precision']*100:.0f}%"] for i in range(5)])
img("ragas_cmp_all_ragas.png")

# Section 4: Analysis
doc.add_page_break()
heading("4. Analysis: Why the Scores Differ")

heading("4.1 Two Factors at Play", level=2)
doc.add_paragraph(
    "The difference between RAGAS and custom scores comes from two factors working together:")
doc.add_paragraph(
    "Judge quality: The custom evaluation used Ollama's 3B model as judge, which is too small "
    "to reliably assess faithfulness and context recall. RAGAS used Groq's 70B model, which "
    "understands the evaluation task much better.", style="List Bullet")
doc.add_paragraph(
    "Evaluation methodology: The custom approach asks for a single holistic score with one prompt. "
    "RAGAS decomposes the evaluation into atomic checks — extracting individual claims and verifying "
    "each one separately. This structured approach is more rigorous and less prone to inconsistency.", style="List Bullet")

heading("4.2 Faithfulness: 92.5% (RAGAS) vs 56.6% (Custom)", level=2)
doc.add_paragraph(
    "The custom evaluation underscored the answers by ~36 percentage points. The 3B judge often "
    "gave low scores even when the answer was clearly grounded in the context. RAGAS's claim-by-claim "
    "verification found that most claims in the answers were actually supported — the answers were "
    "better than the small judge thought.")
doc.add_paragraph(
    "RAGAS scored 92.5% rather than 100% because it found some answers contained claims that went "
    "slightly beyond the context (e.g., the melanoma answer at 91% and diabetes answer at 71%). "
    "This granularity is a strength of the decomposition approach.")

heading("4.3 Context Recall: 100% (RAGAS) vs 38.6% (Custom)", level=2)
doc.add_paragraph(
    "This is the biggest gap. RAGAS found that the retrieved context covered 100% of the reference "
    "information for all questions. The custom evaluation's 38.6% was almost entirely due to the "
    "weak 3B judge — it couldn't properly assess whether the context covered the reference.")

heading("4.4 Context Precision: 93.3% (RAGAS only)", level=2)
doc.add_paragraph(
    "This metric is unique to RAGAS. It measures whether the retrieved context chunks were actually "
    "useful for answering the question. At 93.3%, most retrieved chunks were relevant. "
    "The melanoma and MRSA questions scored 83% — some retrieved chunks about treatment and diagnosis "
    "weren't directly useful for the specific questions asked.")

heading("4.5 BLEU: 51.9% (Custom only)", level=2)
doc.add_paragraph(
    "BLEU is purely algorithmic and not part of RAGAS. It provides an objective baseline that "
    "doesn't depend on any LLM judge. The moderate score reflects that the Ollama answers "
    "captured the right information but used different wording than the reference.")

# Section 5: Key Takeaways
heading("5. Key Takeaways")
doc.add_paragraph(
    "RAGAS provides more reliable and granular evaluation than simple prompt-based scoring. "
    "Its claim decomposition approach catches nuances that a single-prompt judge misses.", style="List Number")
doc.add_paragraph(
    "The 3B Ollama judge severely underscored the answers — by 36% on faithfulness and 61% on "
    "context recall. Small models should not be used as evaluation judges.", style="List Number")
doc.add_paragraph(
    "RAGAS adds Context Precision as a bonus metric, helping identify when the retrieval pipeline "
    "returns irrelevant documents.", style="List Number")
doc.add_paragraph(
    "Combining RAGAS (structured LLM evaluation) with BLEU (algorithmic baseline) gives the most "
    "complete picture of RAG pipeline quality.", style="List Number")
doc.add_paragraph(
    "For production RAG systems, RAGAS is the recommended evaluation framework due to its "
    "rigour, reproducibility, and industry adoption.", style="List Number")

doc.save("RAGAS_vs_Custom_Comparison_Report.docx")
print("Saved: RAGAS_vs_Custom_Comparison_Report.docx")
