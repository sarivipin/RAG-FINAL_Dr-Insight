"""Generate RAGAS chunking comparison chart and Word report."""
import json
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

with open("chunking_ragas_results.json") as f:
    data = json.load(f)

KEYS = ["qa_pair", "recursive", "semantic"]
LABELS = [data[k]["label"] for k in KEYS]
COLORS = ["#4A90D9", "#2ECC71", "#E8833A"]
QS = ["Bacterial\nVaginosis", "Varicose\nVeins", "Melanoma\nPrevention", "Type 1\nDiabetes", "MRSA"]
QF = ["Symptoms of bacterial vaginosis", "Causes of varicose veins",
      "Preventing melanoma", "Living with type 1 diabetes", "MRSA"]

avgs = {k: [data[k]["avg_faithfulness"], data[k]["avg_context_recall"], data[k]["avg_context_precision"]] for k in KEYS}
pq = {k: {m: [r[m] for r in data[k]["per_question"]] for m in ["faithfulness","context_recall","context_precision"]} for k in KEYS}

# =========================================================
# CHARTS
# =========================================================
# Chart 1: Overall
fig, ax = plt.subplots(figsize=(10, 5.5))
metrics = ["Faithfulness", "Context Recall", "Context Precision"]
x = np.arange(len(metrics)); w = 0.22
for i, (k, l, c) in enumerate(zip(KEYS, LABELS, COLORS)):
    vals = [avgs[k][j]*100 for j in range(3)]
    bars = ax.bar(x + (i-1)*w, vals, w, label=l, color=c)
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+1, f"{b.get_height():.0f}%", ha="center", fontsize=8)
ax.set_ylabel("Score (%)"); ax.set_title("Chunking Methods: RAGAS Evaluation Overall", fontsize=13, fontweight="bold")
ax.set_xticks(x); ax.set_xticklabels(metrics); ax.set_ylim(0, 118)
ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
plt.tight_layout(); plt.savefig("chunk_ragas_overall.png", dpi=150); plt.close()

# Per-question charts
for metric, label in [("faithfulness","Faithfulness"),("context_recall","Context Recall"),("context_precision","Context Precision")]:
    fig, ax = plt.subplots(figsize=(11, 5.5))
    x = np.arange(5); w = 0.25
    for i, (k, l, c) in enumerate(zip(KEYS, LABELS, COLORS)):
        ax.bar(x + (i-1)*w, [v*100 for v in pq[k][metric]], w, label=l, color=c)
    ax.set_ylabel("Score (%)"); ax.set_title(f"RAGAS {label} per Question by Chunking Method", fontsize=12, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(QS, fontsize=9); ax.set_ylim(0, 118)
    ax.legend(fontsize=9); ax.grid(axis="y", alpha=0.3)
    plt.tight_layout(); plt.savefig(f"chunk_ragas_{metric}.png", dpi=150); plt.close()

print("Charts generated.")

# =========================================================
# WORD DOCUMENT
# =========================================================
doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)

def heading(t, level=1):
    h = doc.add_heading(t, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

def img(path):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# Title
doc.add_paragraph()
t = doc.add_heading("Chunking Methods — RAGAS Evaluation Report", level=0)
for r in t.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("QA-Pair vs Recursive Character vs Semantic Chunking")
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
doc.add_page_break()

# Section 1: Overview
heading("1. Overview")
doc.add_paragraph(
    "This report compares three chunking strategies evaluated using the RAGAS framework. "
    "All methods use Groq's llama-3.3-70b-versatile for both answer generation and RAGAS evaluation. "
    "The same 5 questions and ground-truth references are used across all methods.")

tbl(["", "QA-Pair (Original)", "Recursive Character", "Semantic Chunking"],
    [["Ingestion File", "ingest.py", "ingest_recursive.py", "ingest_semantic.py"],
     ["Method", "Split by ## Question blocks", "800 chars, 200 overlap", "Embedding similarity splits"],
     ["Chunks Created", "~1,950", "2,801", "4,248"],
     ["Vector Store", "db/chroma_db", "db/chroma_db_recursive", "db/chroma_db_semantic"]])

heading("1.1 RAGAS Metrics", level=2)
tbl(["Metric", "What RAGAS Measures", "How It Works"],
    [["Faithfulness", "Is the answer grounded in context?", "Extracts claims from answer,\nverifies each against context"],
     ["Context Recall", "Does context cover the reference?", "Checks each reference sentence\nagainst retrieved context"],
     ["Context Precision", "Were retrieved chunks useful?", "Checks if each context chunk\nhelped answer the question"]])

# Section 2: Overall Results
doc.add_page_break()
heading("2. Overall Results")

tbl(["Metric", "QA-Pair", "Recursive", "Semantic", "Best"],
    [["Faithfulness", f"{avgs['qa_pair'][0]*100:.1f}%", f"{avgs['recursive'][0]*100:.1f}%", f"{avgs['semantic'][0]*100:.1f}%", "QA-Pair"],
     ["Context Recall", f"{avgs['qa_pair'][1]*100:.1f}%", f"{avgs['recursive'][1]*100:.1f}%", f"{avgs['semantic'][1]*100:.1f}%", "QA-Pair"],
     ["Context Precision", f"{avgs['qa_pair'][2]*100:.1f}%", f"{avgs['recursive'][2]*100:.1f}%", f"{avgs['semantic'][2]*100:.1f}%", "QA-Pair"]])

img("chunk_ragas_overall.png")

# Section 3: Per-Question
doc.add_page_break()
heading("3. Per-Question Breakdown")

for metric, label in [("faithfulness","Faithfulness"),("context_recall","Context Recall"),("context_precision","Context Precision")]:
    heading(f"3.{['faithfulness','context_recall','context_precision'].index(metric)+1} {label}", level=2)
    tbl(["Question", "QA-Pair", "Recursive", "Semantic"],
        [[QF[i], f"{pq['qa_pair'][metric][i]*100:.0f}%", f"{pq['recursive'][metric][i]*100:.0f}%",
          f"{pq['semantic'][metric][i]*100:.0f}%"] for i in range(5)])
    img(f"chunk_ragas_{metric}.png")

# Section 4: Analysis
doc.add_page_break()
heading("4. Analysis")

heading("4.1 QA-Pair Chunking Wins Across All RAGAS Metrics", level=2)
doc.add_paragraph(
    "QA-Pair chunking scored highest on all three RAGAS metrics: 98% faithfulness, 100% context recall, "
    "and 90% context precision. This confirms what our earlier custom evaluation found — when documents "
    "are structured as question-answer pairs, chunking by those natural boundaries preserves the "
    "semantic completeness that retrieval needs.")

heading("4.2 Context Recall: The Biggest Differentiator", level=2)
doc.add_paragraph(
    "Context recall dropped significantly with alternative chunking: 100% (QA-Pair) → 80% (Recursive) → "
    "66% (Semantic). RAGAS verified this by checking each sentence in the reference answer against the "
    "retrieved context. Recursive and semantic splitting fragment QA pairs, so the retriever can't find "
    "complete answers — parts of the reference information are missing from the retrieved chunks.")

heading("4.3 Context Precision: Retrieval Noise", level=2)
doc.add_paragraph(
    "Context precision measures whether retrieved chunks were actually useful. QA-Pair scored 90%, "
    "while Recursive (78.3%) and Semantic (71.7%) scored lower. This means the alternative methods "
    "retrieved more irrelevant chunks — noise that doesn't help answer the question. "
    "More chunks doesn't mean better retrieval.")

heading("4.4 Faithfulness: Similar Across Methods", level=2)
doc.add_paragraph(
    "All three methods scored similarly on faithfulness (97-98%). This makes sense — faithfulness "
    "measures whether the answer stays grounded in whatever context was retrieved. Even with worse "
    "retrieval, the LLM was disciplined about not hallucinating beyond the provided context.")

heading("4.5 RAGAS vs Custom Evaluation Agreement", level=2)
doc.add_paragraph(
    "Both RAGAS and our earlier custom evaluation reached the same conclusion: QA-Pair chunking is best "
    "for this dataset. However, RAGAS provides more granular insights — the context precision metric "
    "revealed retrieval noise that the custom evaluation couldn't detect.")

# Section 5: Conclusion
heading("5. Conclusion")
doc.add_paragraph(
    "QA-Pair chunking is the clear winner for structured FAQ-style medical documents. "
    "RAGAS evaluation confirms this with rigorous, decomposed metrics that go beyond simple scoring.", style="List Number")
doc.add_paragraph(
    "Chunking strategy primarily affects retrieval quality (context recall and precision), "
    "not generation quality (faithfulness). The LLM generates faithful answers regardless — "
    "the bottleneck is whether the right documents are retrieved.", style="List Number")
doc.add_paragraph(
    "RAGAS's context precision metric adds value over custom evaluation by detecting retrieval noise — "
    "an important signal for optimising the retrieval pipeline.", style="List Number")

doc.save("Chunking_RAGAS_Comparison_Report.docx")
print("Saved: Chunking_RAGAS_Comparison_Report.docx")
